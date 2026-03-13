"""
Student Lookup Web App
Search students by name/NetID, view Outlook emails, sync grades from Blackboard.
"""

import os
import json
import sqlite3
import threading
import time
import glob as glob_mod
import urllib.parse as urllib_parse
import requests as http_requests
from flask import Flask, request, redirect, url_for, flash, jsonify, get_flashed_messages

# ── Configuration ────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(SCRIPT_DIR, "students.db")
CHROME_PROFILE = os.path.join(SCRIPT_DIR, "chrome_profile")
DOWNLOAD_DIR = os.path.join(SCRIPT_DIR, "downloads")
PORT = 8081
MAX_EMAILS = 15
BB_URL = "https://elearning.utdallas.edu"

COURSE_LABELS = {
    "buan4320.501": "BUAN 4320.501 – DB Fundamentals (Wed)",
    "buan4320.502": "BUAN 4320.502 – DB Fundamentals (Tue)",
    "buan6320.s01": "BUAN 6320.S01 – DB Foundations (Mon)",
    "itss4351.002": "ITSS 4351.002 – Business Intel (Wed)",
    "itss4351.003": "ITSS 4351.003 – Business Intel (Tue)",
    "buan4351.002": "BUAN 4351.002 – Business Intel (Wed)",
    "buan4351.003": "BUAN 4351.003 – Business Intel (Tue)",
}

# Blackboard course IDs for grade center links
BB_COURSE_IDS = {
    "buan4320.501": "_402886_1",
    "buan4320.502": "_402887_1",
    "buan6320.s01": "_402933_1",
    "itss4351.002": "_410533_1",
    "itss4351.003": "_410592_1",
    "buan4351.002": "_410533_1",
    "buan4351.003": "_410592_1",
}

def bb_gradebook_url(course_raw, student_name=""):
    """Return Blackboard gradebook URL for a course, optionally with student search."""
    ck = course_key(course_raw)
    bb_id = BB_COURSE_IDS.get(ck)
    if not bb_id:
        return ""
    return f"https://elearning.utdallas.edu/ultra/courses/{bb_id}/cl/outline"

app = Flask(__name__)
app.secret_key = "student-lookup-local"

# Global sync status
sync_status = {"running": False, "messages": [], "done": False, "error": None}
rules_status = {"running": False, "messages": [], "done": False, "error": None}

# ── Database helpers ─────────────────────────────────────────────────────────

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS students (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            last_name TEXT NOT NULL,
            first_name TEXT NOT NULL,
            netid TEXT NOT NULL,
            course TEXT NOT NULL,
            UNIQUE(netid, course)
        )
    """)
    # Add grade column if missing
    cols = [r[1] for r in conn.execute("PRAGMA table_info(students)").fetchall()]
    if "grade" not in cols:
        conn.execute("ALTER TABLE students ADD COLUMN grade TEXT DEFAULT ''")
    if "assignments" not in cols:
        conn.execute("ALTER TABLE students ADD COLUMN assignments TEXT DEFAULT ''")
    conn.commit()
    conn.close()

def course_label(raw):
    parts = raw.split(".")
    if len(parts) >= 2:
        key = parts[0].lower() + "." + parts[1].lower()
    else:
        key = raw.lower()
    return COURSE_LABELS.get(key, raw.upper())

def course_key(raw):
    parts = raw.strip().lower().split(".")
    if len(parts) >= 2:
        return parts[0] + "." + parts[1]
    return raw.strip().lower()

# ── Outlook email search ─────────────────────────────────────────────────────

def search_outlook_emails(first_name, last_name):
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    try:
        outlook = win32com.client.Dispatch("Outlook.Application")
        ns = outlook.GetNamespace("MAPI")
        inbox = ns.GetDefaultFolder(6)

        # Collect folders to search: inbox + all "Student Emails" subfolders
        folders_to_search = [inbox]
        try:
            for i in range(1, inbox.Folders.Count + 1):
                if inbox.Folders.Item(i).Name == "Student Emails":
                    parent = inbox.Folders.Item(i)
                    for j in range(1, parent.Folders.Count + 1):
                        folders_to_search.append(parent.Folders.Item(j))
                    break
        except Exception:
            pass

        emails = []
        first_lower = first_name.lower()
        last_lower = last_name.lower()
        dasl = f'@SQL="urn:schemas:httpmail:fromname" LIKE \'%{last_name}%\''

        for folder in folders_to_search:
            if len(emails) >= MAX_EMAILS:
                break
            try:
                items = folder.Items.Restrict(dasl)
                items.Sort("[ReceivedTime]", True)
            except Exception:
                continue

            for i in range(min(MAX_EMAILS * 2, items.Count)):
                if len(emails) >= MAX_EMAILS:
                    break
                item = items.Item(i + 1)
                try:
                    sender_name = item.SenderName or ""
                    name_lower = sender_name.lower()
                    if last_lower not in name_lower:
                        continue
                    if first_lower and first_lower not in name_lower:
                        continue

                    body = (item.Body or "")[:300].strip()
                    body = " ".join(body.split())
                    if len(body) > 200:
                        body = body[:200] + "..."

                    sender_email = ""
                    try:
                        if item.SenderEmailType == "EX":
                            exu = item.Sender.GetExchangeUser()
                            if exu:
                                sender_email = exu.PrimarySmtpAddress
                        else:
                            sender_email = item.SenderEmailAddress
                    except Exception:
                        sender_email = item.SenderEmailAddress or ""

                    received = item.ReceivedTime
                    date_str = received.strftime("%b %d, %Y %I:%M %p")

                    emails.append({
                        "subject": item.Subject or "(no subject)",
                        "date": date_str,
                        "preview": body,
                        "sender": sender_name,
                        "sender_email": sender_email,
                    })
                except Exception:
                    continue

        # Sort all collected emails by date descending
        emails.sort(key=lambda e: e["date"], reverse=True)
        return emails[:MAX_EMAILS]
    finally:
        pythoncom.CoUninitialize()

# ── Blackboard Grade Sync ────────────────────────────────────────────────────

def sync_log(msg):
    sync_status["messages"].append(msg)
    print(f"  [sync] {msg}")

def rules_log(msg):
    rules_status["messages"].append(msg)
    print(f"  [rules] {msg}")

def _open_bb_chrome_and_login():
    """Open Chrome with saved profile, navigate to BB, wait for login. Returns driver."""
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options

    os.makedirs(CHROME_PROFILE, exist_ok=True)
    os.makedirs(DOWNLOAD_DIR, exist_ok=True)

    opts = Options()
    opts.add_argument(f"--user-data-dir={CHROME_PROFILE}")
    opts.add_argument("--profile-directory=Default")
    opts.add_experimental_option("prefs", {
        "download.default_directory": DOWNLOAD_DIR,
        "download.prompt_for_download": False,
    })

    sync_log("Opening Chrome...")
    driver = webdriver.Chrome(options=opts)
    driver.set_window_size(1100, 800)

    sync_log("Navigating to eLearning...")
    driver.get(BB_URL)
    time.sleep(3)

    waited = 0
    while waited < 180:
        try:
            url = (driver.current_url or "").lower()
        except Exception:
            url = ""
        if waited > 0 and waited % 20 == 0:
            sync_log(f"  Current URL: {url[:120]}")
        if "elearning.utdallas.edu" in url and ("ultra" in url or "webapps" in url):
            break
        if "elearning.utdallas.edu" in url and "sso" not in url and "login" not in url and "cas" not in url and "idp" not in url:
            break
        if waited == 0:
            sync_log("Waiting for you to log in (SSO + Duo)...")
        time.sleep(2)
        waited += 2

    if waited >= 180:
        driver.quit()
        raise TimeoutError("Login timed out after 3 minutes.")

    sync_log("Logged in!")
    time.sleep(2)
    return driver


def _download_single_course(driver, bb_id, course_keys):
    """Download and load grade file for a single BB course ID."""
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC

    label = " / ".join(course_keys)

    # Clear previous downloads
    for f in glob_mod.glob(os.path.join(DOWNLOAD_DIR, "*.csv")):
        os.remove(f)
    for f in glob_mod.glob(os.path.join(DOWNLOAD_DIR, "*.xls")):
        os.remove(f)
    for f in glob_mod.glob(os.path.join(DOWNLOAD_DIR, "*.xlsx")):
        os.remove(f)

    # Navigate to classic grade download page
    download_url = (f"{BB_URL}/webapps/gradebook/do/instructor/"
                  f"downloadGradebook?dispatch=viewDownloadOptions"
                  f"&course_id={bb_id}")
    driver.get(download_url)
    time.sleep(4)

    try:
        submit = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, "input[type='submit'], button[type='submit']")))
        submit.click()
        sync_log(f"  Download triggered, waiting for file...")
    except Exception:
        sync_log(f"  Classic download page not available, trying Ultra gradebook...")
        driver.get(f"{BB_URL}/ultra/courses/{bb_id}/gradebook")
        time.sleep(3)
        try:
            menu_btns = driver.find_elements(By.CSS_SELECTOR, "[data-testid='kebab-menu'], .kebab-menu, button[aria-label*='ownload']")
            for btn in menu_btns:
                btn.click()
                time.sleep(1)
            download_btns = driver.find_elements(By.XPATH, "//*[contains(text(),'Download') or contains(text(),'Export')]")
            for btn in download_btns:
                btn.click()
                time.sleep(2)
                break
        except Exception as e2:
            raise RuntimeError(f"Could not trigger download: {e2}")

    # Wait for file to appear
    waited = 0
    downloaded_file = None
    while waited < 30:
        time.sleep(2)
        waited += 2
        files = glob_mod.glob(os.path.join(DOWNLOAD_DIR, "*.csv"))
        files += glob_mod.glob(os.path.join(DOWNLOAD_DIR, "*.xls"))
        files += glob_mod.glob(os.path.join(DOWNLOAD_DIR, "*.xlsx"))
        files = [f for f in files if not f.endswith(".crdownload")]
        if files:
            downloaded_file = max(files, key=os.path.getmtime)
            break

    if not downloaded_file:
        raise RuntimeError("No file downloaded after 30s")

    sync_log(f"  Downloaded: {os.path.basename(downloaded_file)}")

    with open(downloaded_file, "rb") as fobj:
        class FileWrapper:
            def __init__(self, data, name):
                self.data = data
                self.filename = name
            def read(self):
                return self.data
        fw = FileWrapper(fobj.read(), os.path.basename(downloaded_file))
        count, courses = load_master_grade(fw)
        sync_log(f"  Loaded {count} students with grades")
    return count


def run_single_course_sync(bb_id, course_keys):
    """Sync grades for a single course — opens Chrome, downloads, loads."""
    global sync_status
    sync_status = {"running": True, "messages": [], "done": False, "error": None}
    label = " / ".join(course_keys)
    try:
        driver = _open_bb_chrome_and_login()
        try:
            sync_log(f"Syncing {label}...")
            _download_single_course(driver, bb_id, course_keys)
            sync_log("Done!")
        finally:
            sync_log("Closing browser...")
            driver.quit()
    except Exception as e:
        sync_status["error"] = str(e)
        sync_log(f"Error: {e}")
    finally:
        sync_status["running"] = False
        sync_status["done"] = True


def run_grade_sync():
    """Open Chrome, login to Blackboard, download grades for all SP26 courses."""
    global sync_status
    sync_status = {"running": True, "messages": [], "done": False, "error": None}

    try:
        driver = _open_bb_chrome_and_login()
        try:
            sync_log(f"Downloading grades for Spring 2026 courses...")
            grade_sync_via_download_whitelist(driver)
        finally:
            sync_log("Closing browser...")
            driver.quit()

    except Exception as e:
        sync_status["error"] = str(e)
        sync_log(f"Error: {e}")
    finally:
        sync_status["running"] = False
        sync_status["done"] = True


def grade_sync_via_download_whitelist(driver):
    """Download Full Grade Center for each Spring 2026 course using Selenium."""
    seen_bb_ids = {}
    for course_key_str, bb_id in BB_COURSE_IDS.items():
        if bb_id not in seen_bb_ids:
            seen_bb_ids[bb_id] = []
        seen_bb_ids[bb_id].append(course_key_str)

    sync_log(f"Downloading grades for {len(seen_bb_ids)} courses...")
    for bb_id, course_keys in seen_bb_ids.items():
        label = " / ".join(course_keys)
        sync_log(f"Downloading: {label} (BB ID: {bb_id})...")
        try:
            _download_single_course(driver, bb_id, course_keys)
        except Exception as e:
            sync_log(f"  Error: {e}")
    sync_log("Done!")


def _grade_sync_via_api_direct_UNUSED(session, api):
    """OLD: Fetch grades using BB_COURSE_IDS whitelist — only Spring 2026 courses."""
    # Deduplicate: multiple course keys can map to same bb_id (merged courses)
    seen_bb_ids = {}
    for course_key_str, bb_id in BB_COURSE_IDS.items():
        if bb_id not in seen_bb_ids:
            seen_bb_ids[bb_id] = []
        seen_bb_ids[bb_id].append(course_key_str)

    synced = 0
    for bb_id, course_keys in seen_bb_ids.items():
        label = " / ".join(course_keys)
        sync_log(f"Syncing {label} (BB ID: {bb_id})...")

        # Get gradebook columns
        try:
            cols_resp = session.get(f"{api}/courses/{bb_id}/gradebook/columns",
                                   params={"limit": 200}, timeout=15)
            if cols_resp.status_code != 200:
                sync_log(f"  Could not fetch columns (HTTP {cols_resp.status_code})")
                continue

            columns = cols_resp.json().get("results", [])
            if not columns:
                sync_log(f"  No grade columns found")
                continue

            # Separate total column from assignment columns
            total_col = None
            assignment_cols = []
            for col in columns:
                col_name = col.get("name") or ""
                col_name_lower = col_name.lower()
                if "total" in col_name_lower or "weighted" in col_name_lower or "final" in col_name_lower:
                    # Prefer non-weighted total; keep first total found
                    if total_col is None or ("weighted" not in col_name_lower and "total" in col_name_lower):
                        total_col = col
                else:
                    # It's an individual assignment column
                    score_obj = col.get("score") or {}
                    max_pts = str(score_obj.get("possible") or "")
                    if not max_pts:
                        max_pts = str(col.get("possible") or "")
                    assignment_cols.append({
                        "id": col["id"],
                        "name": col_name,
                        "max": max_pts,
                    })

            if not total_col and columns:
                total_col = columns[-1]  # Last column as fallback

            sync_log(f"  Found {len(assignment_cols)} assignment columns + total column")

            # Get user list to map user IDs to usernames (netids)
            user_map = {}
            users_resp = session.get(f"{api}/courses/{bb_id}/users",
                                     params={"limit": 500}, timeout=15)
            if users_resp.status_code == 200:
                users_data = users_resp.json().get("results", [])
                sync_log(f"  Found {len(users_data)} enrolled users")
                if users_data:
                    sync_log(f"  DEBUG first user: {json.dumps(users_data[0])[:500]}")
                for u in users_data:
                    uid = u.get("userId") or ""
                    uname = u.get("userName") or ""
                    if not uname:
                        user_obj = u.get("user") or {}
                        uname = user_obj.get("userName") or ""
                    if uid and uname:
                        user_map[uid] = uname
                # If membership data has no userName, look up each user by ID
                if not user_map and users_data:
                    sync_log(f"  No userNames in membership, looking up users individually...")
                    for u in users_data:
                        uid = u.get("userId") or ""
                        if not uid:
                            continue
                        try:
                            ur = session.get(f"{api}/users/{uid}", timeout=10)
                            if ur.status_code == 200:
                                ud = ur.json()
                                uname = ud.get("userName") or ""
                                if uid and uname:
                                    user_map[uid] = uname
                        except Exception:
                            continue
                    if user_map:
                        sync_log(f"  DEBUG sample user lookup: {json.dumps(ud)[:500]}")
            else:
                sync_log(f"  Could not fetch users (HTTP {users_resp.status_code})")
            sync_log(f"  Mapped {len(user_map)} users to netids")

            # Build per-user assignment scores: {userId: {col_id: score_str}}
            user_scores = {}

            # Fetch grades for each assignment column
            for acol in assignment_cols:
                col_id = acol["id"]
                try:
                    gr = session.get(
                        f"{api}/courses/{bb_id}/gradebook/columns/{col_id}/users",
                        params={"limit": 200}, timeout=15)
                    if gr.status_code != 200:
                        continue
                    for g in gr.json().get("results", []):
                        uid = g.get("userId", "")
                        score = g.get("displayGrade", {}).get("text", "")
                        if not score:
                            score = g.get("score", "")
                        if score is not None:
                            score = str(score).strip()
                        else:
                            score = ""
                        user_scores.setdefault(uid, {})[col_id] = score
                except Exception:
                    continue

            # Fetch total column grades
            total_grades = {}
            if total_col:
                try:
                    tr = session.get(
                        f"{api}/courses/{bb_id}/gradebook/columns/{total_col['id']}/users",
                        params={"limit": 200}, timeout=15)
                    if tr.status_code == 200:
                        for g in tr.json().get("results", []):
                            uid = g.get("userId", "")
                            score = g.get("displayGrade", {}).get("text", "")
                            if not score:
                                score = str(g.get("score", ""))
                            total_grades[uid] = score
                except Exception:
                    pass

            # Update DB with total grade + assignments JSON
            conn = get_db()
            updated = 0
            for uid, netid in user_map.items():
                if not netid:
                    continue

                # Total grade
                grade_str = total_grades.get(uid, "")
                if grade_str:
                    try:
                        score_num = float(grade_str)
                        grade_str = f"{score_num:.1f}" if score_num != int(score_num) else str(int(score_num))
                    except (ValueError, TypeError):
                        pass

                # Assignments JSON
                assignments = []
                scores = user_scores.get(uid, {})
                for acol in assignment_cols:
                    score_str = scores.get(acol["id"], "")
                    if score_str:
                        try:
                            sv = float(score_str)
                            score_str = f"{sv:.1f}" if sv != int(sv) else str(int(sv))
                        except (ValueError, TypeError):
                            pass
                    assignments.append({
                        "name": acol["name"],
                        "score": score_str,
                        "max": acol["max"],
                    })

                assignments_json = json.dumps(assignments) if assignments else ""

                for ck in course_keys:
                    conn.execute(
                        "UPDATE students SET grade = ?, assignments = ? WHERE netid = ? AND course = ?",
                        (grade_str, assignments_json, netid, ck))
                updated += 1

            conn.commit()
            conn.close()
            sync_log(f"  Updated {updated} students ({len(assignment_cols)} assignments each)")
            synced += 1

        except Exception as e:
            sync_log(f"  Error fetching grades: {e}")

    sync_log(f"Done! Synced {synced}/{len(seen_bb_ids)} courses.")


def grade_sync_via_download(driver, session):
    """Fallback: navigate to grade center and download files."""
    sync_log("Navigating to Blackboard courses page...")

    # Try to get course list from the page
    driver.get(f"{BB_URL}/ultra/institution-page")
    time.sleep(3)

    sync_log("Looking for courses...")
    # Try to find course links
    try:
        from selenium.webdriver.common.by import By
        links = driver.find_elements(By.CSS_SELECTOR, "a[href*='/ultra/courses/']")
        course_links = []
        for link in links:
            href = link.get_attribute("href") or ""
            text = link.text.strip()
            if "/ultra/courses/" in href and text:
                course_links.append({"href": href, "text": text})

        if not course_links:
            sync_log("Could not find course links on the page.")
            sync_log("Please download grade files manually from Blackboard and upload them.")
            return

        sync_log(f"Found {len(course_links)} courses on dashboard")

        for cl in course_links:
            sync_log(f"  - {cl['text']}")

        # For each course, try to navigate to gradebook and download
        for cl in course_links:
            course_url = cl["href"]
            # Extract course ID from URL
            parts = course_url.split("/ultra/courses/")
            if len(parts) < 2:
                continue
            bb_course_id = parts[1].split("/")[0]

            sync_log(f"Downloading grades for: {cl['text']}...")

            # Try the classic grade download URL
            download_url = (f"{BB_URL}/webapps/gradebook/do/instructor/"
                          f"downloadGradebook?dispatch=viewDownloadOptions"
                          f"&course_id={bb_course_id}")
            driver.get(download_url)
            time.sleep(3)

            # Look for download button and click
            try:
                submit = driver.find_element(By.CSS_SELECTOR, "input[type='submit']")
                submit.click()
                time.sleep(5)
                sync_log(f"  Download triggered for {cl['text']}")
            except Exception:
                sync_log(f"  Could not auto-download, page may require manual interaction")

        # Parse any downloaded files
        parse_downloaded_grades()

    except Exception as e:
        sync_log(f"Error during download: {e}")
        sync_log("Please download grade files manually and upload them.")


def parse_downloaded_grades():
    """Parse grade CSV/Excel files from the download directory."""
    import csv

    files = glob_mod.glob(os.path.join(DOWNLOAD_DIR, "*.csv"))
    files += glob_mod.glob(os.path.join(DOWNLOAD_DIR, "*.xlsx"))

    if not files:
        sync_log("No grade files found in download directory.")
        return

    conn = get_db()
    for filepath in files:
        try:
            if filepath.endswith(".csv"):
                with open(filepath, "r", encoding="utf-8-sig") as f:
                    reader = csv.DictReader(f)
                    updated = 0
                    for row in reader:
                        username = row.get("Username", row.get("username", "")).strip()
                        # Look for a "Total" or grade column
                        grade = ""
                        for key in row:
                            if "total" in key.lower() or "grade" in key.lower() or "weighted" in key.lower():
                                grade = row[key].strip()
                                break
                        if username and grade:
                            conn.execute("UPDATE students SET grade = ? WHERE netid = ?",
                                         (grade, username))
                            updated += 1
                    sync_log(f"  Parsed {updated} grades from {os.path.basename(filepath)}")
            # Could add .xlsx parsing here too
        except Exception as e:
            sync_log(f"  Error parsing {os.path.basename(filepath)}: {e}")

    conn.commit()
    conn.close()


# ── Outlook Email Sorter ────────────────────────────────────────────────────

# Background scanner state
_scanner_thread = None
SCAN_INTERVAL = 300  # seconds (5 minutes)


def _course_folder_name(course_raw):
    """Short folder name from course key: 'buan4320.501' → 'BUAN 4320.501'."""
    import re
    m = re.match(r'([a-z]+)(\d+)\.(.+)', course_raw, re.IGNORECASE)
    if m:
        return f"{m.group(1).upper()} {m.group(2)}.{m.group(3).upper()}"
    return course_raw.upper()


def _find_course_folder(parent, course_raw):
    """Find an existing subfolder matching this course (by prefix). Returns folder or None."""
    short = _course_folder_name(course_raw)
    for i in range(1, parent.Folders.Count + 1):
        fname = parent.Folders.Item(i).Name.strip()
        if fname == short or fname.startswith(short):
            return parent.Folders.Item(i)
    return None


def run_setup_outlook_rules():
    """Thread target: create folders, move existing emails, start background scanner."""
    global rules_status
    rules_status = {"running": True, "messages": [], "done": False, "error": None}

    try:
        import pythoncom
        import win32com.client

        pythoncom.CoInitialize()
        try:
            rules_log("Connecting to Outlook...")
            outlook = win32com.client.Dispatch("Outlook.Application")
            ns = outlook.GetNamespace("MAPI")
            inbox = ns.GetDefaultFolder(6)  # olFolderInbox

            # Get all courses + student netids from DB
            conn = get_db()
            courses = conn.execute("""
                SELECT course, GROUP_CONCAT(netid) as netids, COUNT(*) as cnt
                FROM students GROUP BY course ORDER BY course
            """).fetchall()
            conn.close()

            if not courses:
                rules_status["error"] = "No students in database. Upload rosters first."
                return

            rules_log(f"Found {len(courses)} courses in database")

            # Create "Student Emails" parent folder under Inbox
            parent_folder = None
            for i in range(1, inbox.Folders.Count + 1):
                if inbox.Folders.Item(i).Name == "Student Emails":
                    parent_folder = inbox.Folders.Item(i)
                    break
            if parent_folder is None:
                parent_folder = inbox.Folders.Add("Student Emails")
                rules_log("Created 'Student Emails' folder in Inbox")
            else:
                rules_log("Found existing 'Student Emails' folder")

            # Create course subfolders and build folder mapping
            course_folders = {}
            total_students = 0

            for c in courses:
                course_raw = c['course']
                folder_name = _course_folder_name(course_raw)
                netids = [n.strip() for n in c['netids'].split(',')]
                total_students += len(netids)

                # Find existing folder (by prefix) or create new one
                course_folder = _find_course_folder(parent_folder, course_raw)
                if course_folder is None:
                    course_folder = parent_folder.Folders.Add(folder_name)

                course_folders[course_raw] = course_folder
                rules_log(f"Folder ready: {course_folder.Name} ({len(netids)} students)")

            # Build name-based lookup for matching senders to students
            by_last_name, by_netid = _build_student_lookup()
            rules_log(f"Matching by name + netid ({len(by_last_name)} last names, {len(by_netid)} netids)")

            # Move existing inbox emails to course folders
            rules_log("Scanning inbox for existing student emails...")
            total_moved = _move_existing_student_emails(
                inbox, course_folders, by_last_name, by_netid)
            if total_moved > 0:
                rules_log(f"Moved {total_moved} existing emails to course folders")
            else:
                rules_log("No student emails found in inbox")

            rules_log(f"Done! {len(courses)} folders, {total_students} students")

        finally:
            pythoncom.CoUninitialize()

        # Start background scanner
        _start_scanner()
        rules_log(f"Background scanner started (checks every {SCAN_INTERVAL // 60} min)")

    except Exception as e:
        rules_status["error"] = str(e)
        rules_log(f"Error: {e}")
    finally:
        rules_status["running"] = False
        rules_status["done"] = True


def _get_sender_info(item):
    """Extract SMTP email and display name from an Outlook mail item."""
    sender_name = (item.SenderName or "").strip()
    try:
        if item.SenderEmailType == "EX":
            exu = item.Sender.GetExchangeUser()
            if exu:
                return (exu.PrimarySmtpAddress or "").lower(), sender_name
        return (item.SenderEmailAddress or "").lower(), sender_name
    except Exception:
        return (getattr(item, 'SenderEmailAddress', '') or "").lower(), sender_name


def _build_student_lookup():
    """Build lookup tables for matching inbox senders to student courses.

    Returns (by_last_name, by_netid):
      by_last_name: {last_lower: [(first_lower, course)]}
      by_netid: {netid_lower: course}
    """
    conn = get_db()
    students = conn.execute(
        "SELECT last_name, first_name, netid, course FROM students"
    ).fetchall()
    conn.close()

    by_last_name = {}
    by_netid = {}
    for s in students:
        netid = s['netid'].strip().lower()
        ln = s['last_name'].strip().lower()
        fn = s['first_name'].strip().lower()
        course = s['course']
        by_netid[netid] = course
        by_last_name.setdefault(ln, []).append((fn, course))

    return by_last_name, by_netid


def _match_sender_to_course(sender_email, sender_name, by_last_name, by_netid):
    """Match an inbox sender to a student course key. Returns course or None."""
    import re

    # Try 1: netid email match
    if '@utdallas.edu' in sender_email:
        local = sender_email.split('@')[0]
        if local in by_netid:
            return by_netid[local]

    # Try 2: name-based email (firstname.lastname@utdallas.edu)
    if '@utdallas.edu' in sender_email:
        local = sender_email.split('@')[0]
        parts = local.split('.')
        if len(parts) >= 2:
            first_part = parts[0]
            last_part = re.sub(r'\d+$', '', parts[-1])  # strip trailing digits
            if last_part in by_last_name:
                for fn, course in by_last_name[last_part]:
                    fn_first = fn.split()[0]  # first word of first name
                    if fn_first.startswith(first_part) or first_part.startswith(fn_first):
                        return course

    # Try 3: sender display name ("Last, First" format from Exchange)
    if ',' in sender_name:
        parts = sender_name.split(',', 1)
        last = parts[0].strip().lower()
        first = parts[1].strip().split()[0].lower()
        if last in by_last_name:
            for fn, course in by_last_name[last]:
                fn_first = fn.split()[0]
                if fn_first.startswith(first) or first.startswith(fn_first):
                    return course

    return None


def _move_existing_student_emails(inbox, course_folders, by_last_name, by_netid):
    """Move existing inbox emails from students to their course folders."""
    moved = 0
    try:
        dasl = '@SQL="urn:schemas:httpmail:datereceived" >= \'01/01/2026\''
        items = inbox.Items.Restrict(dasl)
        items.Sort("[ReceivedTime]", True)

        count = items.Count
        rules_log(f"  Scanning {count} recent inbox items...")

        for i in range(count, 0, -1):
            try:
                item = items.Item(i)
                sender_email, sender_name = _get_sender_info(item)
                course = _match_sender_to_course(
                    sender_email, sender_name, by_last_name, by_netid)
                if course and course in course_folders:
                    item.Move(course_folders[course])
                    moved += 1
            except Exception:
                continue
    except Exception as e:
        rules_log(f"  Warning scanning inbox: {e}")

    return moved


def _start_scanner():
    """Start (or restart) the background inbox scanner thread."""
    global _scanner_thread
    if _scanner_thread and _scanner_thread.is_alive():
        return  # Already running
    _scanner_thread = threading.Thread(target=_scanner_loop, daemon=True)
    _scanner_thread.start()


def _scanner_loop():
    """Background loop: periodically scan inbox and move student emails."""
    while True:
        time.sleep(SCAN_INTERVAL)
        try:
            moved = _scan_and_move()
            if moved > 0:
                print(f"  [scanner] Moved {moved} student emails")
        except Exception as e:
            print(f"  [scanner] Error: {e}")


def _scan_and_move():
    """One-shot: scan inbox for student emails and move to course folders."""
    import pythoncom
    import win32com.client
    import datetime

    by_last_name, by_netid = _build_student_lookup()
    if not by_netid:
        return 0

    pythoncom.CoInitialize()
    try:
        outlook = win32com.client.Dispatch("Outlook.Application")
        ns = outlook.GetNamespace("MAPI")
        inbox = ns.GetDefaultFolder(6)

        # Find "Student Emails" parent folder
        parent_folder = None
        for i in range(1, inbox.Folders.Count + 1):
            if inbox.Folders.Item(i).Name == "Student Emails":
                parent_folder = inbox.Folders.Item(i)
                break
        if parent_folder is None:
            return 0

        # Build course_key → folder mapping
        course_folders = {}
        conn = get_db()
        courses = conn.execute("SELECT DISTINCT course FROM students").fetchall()
        conn.close()
        for c in courses:
            f = _find_course_folder(parent_folder, c['course'])
            if f:
                course_folders[c['course']] = f

        # Scan recent inbox items (last 7 days)
        cutoff = (datetime.datetime.now() - datetime.timedelta(days=7)).strftime("%m/%d/%Y")
        dasl = f'@SQL="urn:schemas:httpmail:datereceived" >= \'{cutoff}\''
        items = inbox.Items.Restrict(dasl)

        moved = 0
        for i in range(items.Count, 0, -1):
            try:
                item = items.Item(i)
                sender_email, sender_name = _get_sender_info(item)
                course = _match_sender_to_course(
                    sender_email, sender_name, by_last_name, by_netid)
                if course and course in course_folders:
                    item.Move(course_folders[course])
                    moved += 1
            except Exception:
                continue

        return moved
    finally:
        pythoncom.CoUninitialize()


# ── HTML helpers ─────────────────────────────────────────────────────────────

CSS = """
* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: 'Segoe UI', system-ui, -apple-system, sans-serif;
       background: #f5f5f5; color: #333; min-height: 100vh; }
.navbar { background: #c75b12; padding: 0 2rem; display: flex;
           align-items: center; height: 56px; box-shadow: 0 2px 4px rgba(0,0,0,0.15); }
.navbar .brand { color: #fff; font-weight: 700; font-size: 1.2rem;
                  text-decoration: none; margin-right: 2rem; }
.nav-link { color: rgba(255,255,255,0.85); text-decoration: none;
             padding: 0.5rem 1rem; border-radius: 4px; font-size: 0.9rem; }
.nav-link:hover, .nav-link.active { background: rgba(255,255,255,0.15); color: #fff; }
.container { max-width: 960px; margin: 2rem auto; padding: 0 1rem; }
.card { background: #fff; border-radius: 8px; padding: 2rem;
         box-shadow: 0 1px 3px rgba(0,0,0,0.1); margin-bottom: 1.5rem; }
h1 { color: #c75b12; margin-bottom: 1rem; font-size: 1.5rem; }
h2 { color: #008542; margin-bottom: 0.75rem; font-size: 1.2rem; }
input[type=text] { width: 100%; padding: 0.7rem 1rem; border: 2px solid #ddd;
    border-radius: 6px; font-size: 1rem; transition: border-color 0.2s; }
input[type=text]:focus { outline: none; border-color: #c75b12; }
.btn { background: #c75b12; color: #fff; border: none; padding: 0.7rem 1.5rem;
        border-radius: 6px; font-size: 1rem; cursor: pointer; font-weight: 600;
        text-decoration: none; display: inline-block; }
.btn:hover { background: #a94a0e; }
.btn-green { background: #008542; }
.btn-green:hover { background: #006b35; }
.btn-sm { padding: 0.3rem 0.7rem; font-size: 0.8rem; border-radius: 4px; }
table { width: 100%; border-collapse: collapse; margin-top: 1rem; }
th { background: #008542; color: #fff; padding: 0.6rem 0.8rem;
     text-align: left; font-size: 0.85rem; text-transform: uppercase; letter-spacing: 0.5px; }
td { padding: 0.6rem 0.8rem; border-bottom: 1px solid #eee; font-size: 0.9rem; vertical-align: top; }
tr:hover > td { background: #f9f9f9; }
a { color: #c75b12; text-decoration: none; }
a:hover { text-decoration: underline; }
.flash { padding: 0.8rem 1rem; border-radius: 6px; margin-bottom: 1rem; font-size: 0.9rem; }
.flash-success { background: #d4edda; color: #155724; border: 1px solid #c3e6cb; }
.flash-error { background: #f8d7da; color: #721c24; border: 1px solid #f5c6cb; }
.search-box { display: flex; gap: 0.5rem; }
.search-box input { flex: 1; }
.stats { color: #666; font-size: 0.85rem; margin-top: 0.5rem; }
.ng-grid { display: grid; grid-template-columns: repeat(auto-fill, minmax(260px, 1fr)); gap: 0.75rem; margin-top: 0.75rem; }
.ng-card { background: #fff; border: 1px solid #e0e0e0; border-radius: 8px; padding: 0.75rem 1rem;
           display: flex; justify-content: space-between; align-items: center; }
.ng-card:hover { box-shadow: 0 2px 6px rgba(0,0,0,0.1); }
.ng-course { font-weight: 600; font-size: 0.85rem; color: #333; }
.ng-badge { display: inline-flex; align-items: center; gap: 0.3rem; font-size: 0.8rem; font-weight: 700; }
.ng-count { background: #f44336; color: #fff; padding: 0.15rem 0.5rem; border-radius: 10px; font-size: 0.78rem; }
.ng-zero { background: #4caf50; color: #fff; padding: 0.15rem 0.5rem; border-radius: 10px; font-size: 0.78rem; }
.ng-student-count { font-size: 0.78rem; color: #666; margin-top: 0.2rem; }
.ng-card a { text-decoration: none; color: inherit; display: flex; justify-content: space-between; align-items: center; width: 100%; }
.ng-tag { display: inline-block; background: #fff3e0; color: #e65100; padding: 0.1rem 0.4rem;
           border-radius: 4px; font-size: 0.72rem; font-weight: 600; margin: 0.1rem 0.15rem 0.1rem 0; }
.ng-none { color: #999; font-size: 0.78rem; font-style: italic; }
.back-link { display: inline-block; margin-bottom: 1rem; font-size: 0.9rem; color: #c75b12; }
.back-link:hover { text-decoration: underline; }
.course-header { display: flex; justify-content: space-between; align-items: center; flex-wrap: wrap; gap: 0.5rem; }
.course-summary { display: flex; gap: 1.5rem; margin: 0.75rem 0 0.5rem 0; flex-wrap: wrap; }
.course-stat { font-size: 0.85rem; color: #555; }
.course-stat strong { color: #333; }
.student-table { width: 100%; border-collapse: collapse; margin-top: 0.75rem; }
.student-table th { background: #008542; color: #fff; padding: 0.5rem 0.7rem;
    text-align: left; font-size: 0.8rem; text-transform: uppercase; letter-spacing: 0.4px; }
.student-table td { padding: 0.45rem 0.7rem; border-bottom: 1px solid #eee; font-size: 0.88rem; }
.student-table tr:hover > td { background: #f6f9f6; }
.ng-detail { font-size: 0.75rem; color: #888; margin-top: 0.25rem; }
.ng-items { list-style: none; padding: 0; margin: 0.3rem 0 0 0; }
.ng-items li { font-size: 0.78rem; color: #666; padding: 0.1rem 0; }
.ng-items li span { color: #c75b12; font-weight: 600; }
.drop-zone { border: 2px dashed #ccc; border-radius: 8px; padding: 3rem 2rem;
              text-align: center; color: #888; transition: all 0.2s; cursor: pointer; }
.drop-zone.dragover { border-color: #c75b12; background: #fff5f0; color: #c75b12; }
.drop-zone p { margin: 0.5rem 0; }
.course-tag { display: inline-block; background: #e8f5e9; color: #2e7d32;
               padding: 0.15rem 0.5rem; border-radius: 4px; font-size: 0.8rem; font-weight: 600; }
.grade-tag { display: inline-block; padding: 0.15rem 0.5rem; border-radius: 4px;
              font-size: 0.85rem; font-weight: 700; }
.grade-A { background: #d4edda; color: #155724; }
.grade-B { background: #cce5ff; color: #004085; }
.grade-C { background: #fff3cd; color: #856404; }
.grade-D { background: #f8d7da; color: #721c24; }
.grade-F { background: #f5c6cb; color: #721c24; }
.grade-default { background: #e2e3e5; color: #383d41; }
.empty { text-align: center; color: #888; padding: 2rem; }
.grades-panel { display: none; }
.grades-panel.open { display: table-row; }
.grades-panel td { padding: 0; background: #f5f9ff; }
.grades-container { padding: 1rem 1.5rem; }
.grades-loading { color: #888; padding: 1rem; font-style: italic; }
.grades-none { color: #999; font-style: italic; padding: 0.5rem 0; }
.grades-header { display: flex; justify-content: space-between; align-items: center; margin-bottom: 0.5rem; }
.grades-header h3 { color: #e65100; font-size: 0.95rem; margin: 0; }
.assign-table { width: 100%; border-collapse: collapse; margin: 0; font-size: 0.85rem; }
.assign-table th { background: #e65100; color: #fff; padding: 0.4rem 0.6rem;
    font-size: 0.78rem; text-transform: uppercase; letter-spacing: 0.3px; }
.assign-table td { padding: 0.35rem 0.6rem; border-bottom: 1px solid #e0e0e0; }
.assign-table tr:last-child td { border-bottom: none; font-weight: 700; background: #eef2f7; }
.assign-table tr:hover td { background: #eef2f7; }
.score-high { color: #155724; font-weight: 600; }
.score-mid { color: #856404; font-weight: 600; }
.score-low { color: #721c24; font-weight: 600; }
.score-none { color: #aaa; }
.btn-blue { background: #e65100; }
.btn-blue:hover { background: #bf4400; }
.btn-sync { background: #1565c0; }
.btn-sync:hover { background: #0d47a1; }
.email-panel { display: none; }
.email-panel.open { display: table-row; }
.email-panel td { padding: 0; background: #fafafa; }
.email-container { padding: 1rem 1.5rem; }
.email-loading { color: #888; padding: 1rem; font-style: italic; }
.email-list { list-style: none; padding: 0; margin: 0; }
.email-item { padding: 0.7rem 0; border-bottom: 1px solid #eee; }
.email-item:last-child { border-bottom: none; }
.email-subject { font-weight: 600; color: #333; font-size: 0.9rem; }
.email-date { color: #888; font-size: 0.78rem; margin-left: 0.5rem; }
.email-preview { color: #666; font-size: 0.82rem; margin-top: 0.25rem; line-height: 1.4; }
.email-none { color: #999; font-style: italic; padding: 0.5rem 0; }
.email-header { display: flex; justify-content: space-between; align-items: center; margin-bottom: 0.5rem; }
.email-header h3 { color: #008542; font-size: 0.95rem; margin: 0; }
.email-count { background: #c75b12; color: #fff; font-size: 0.75rem;
               padding: 0.1rem 0.5rem; border-radius: 10px; font-weight: 600; }
.course-btn { display: inline-block; padding: 0.8rem 1.5rem; margin: 0.5rem;
              border-radius: 8px; font-size: 1rem; font-weight: 700; cursor: pointer;
              border: 2px solid transparent; text-decoration: none; transition: all 0.2s; }
.course-btn:hover { transform: translateY(-2px); box-shadow: 0 4px 8px rgba(0,0,0,0.15); text-decoration: none; }
.course-btn-active { border-color: #c75b12; box-shadow: 0 0 0 3px rgba(199,91,18,0.2); }
.btn-buan6320 { background: #004085; color: #fff; }
.btn-buan6320:hover { background: #003060; }
.btn-buan4320-s01 { background: #c75b12; color: #fff; }
.btn-buan4320-s01:hover { background: #a54a0e; }
.btn-buan4320-s02 { background: #2e7d32; color: #fff; }
.btn-buan4320-s02:hover { background: #1b5e20; }
.btn-buan4320 { background: #008542; color: #fff; }
.btn-buan4320:hover { background: #006b35; }
.btn-itss4351 { background: #6f42c1; color: #fff; }
.btn-itss4351:hover { background: #5a32a3; }
.assign-btn { display: inline-block; padding: 0.5rem 1.2rem; margin: 0.3rem;
              border-radius: 6px; font-size: 0.9rem; font-weight: 600; cursor: pointer;
              background: #e9ecef; color: #333; border: 2px solid #dee2e6;
              text-decoration: none; transition: all 0.2s; }
.assign-btn:hover { background: #dee2e6; text-decoration: none; }
.assign-btn-active { background: #c75b12; color: #fff; border-color: #c75b12; }
.grade-result { margin-top: 1.5rem; }
.grade-result table { width: 100%; }
.grade-result th { background: #004085; }
.comment-cell { font-size: 0.82rem; color: #721c24; max-width: 400px; }
.score-cell { font-weight: 700; font-size: 1.1rem; text-align: center; }
.score-perfect { color: #155724; }
.score-good { color: #004085; }
.score-deducted { color: #c75b12; }
.grading-drop { border: 2px dashed #ccc; border-radius: 8px; padding: 2rem;
                text-align: center; color: #888; transition: all 0.2s; cursor: pointer; margin-top: 1rem; }
.grading-drop.dragover { border-color: #c75b12; background: #fff5f0; color: #c75b12; }
.grading-status { margin-top: 1rem; }
.grading-progress { background: #1e1e1e; color: #d4d4d4; border-radius: 6px; padding: 1rem;
                    font-family: 'Consolas', monospace; font-size: 0.82rem;
                    max-height: 300px; overflow-y: auto; line-height: 1.6; display: none; }
.grading-progress .msg::before { content: "> "; color: #c75b12; }
.toggle-btn { background: none; border: none; cursor: pointer; padding: 0.2rem 0.4rem; font-size: inherit; color: inherit; }
.file-table { width: 100%; border-collapse: collapse; font-size: 0.85rem; }
.file-table th { background: #e65100; color: #fff; padding: 0.4rem 0.6rem; text-align: left; }
.file-table td { padding: 0.4rem 0.6rem; border-bottom: 1px solid #e0e0e0; }
.file-table tr:hover td { background: #fff5f0; }
.file-link { display: inline-block; margin: 0.15rem 0.3rem 0.15rem 0; padding: 0.2rem 0.5rem;
             background: #fff3e0; border: 1px solid #ffe0b2; border-radius: 4px; color: #e65100;
             text-decoration: none; font-size: 0.82rem; }
.file-link:hover { background: #ffe0b2; }
.toggle-btn:hover { background: #f0f0f0; border-radius: 4px; }
.toggle-arrow { font-size: 0.7rem; color: #888; margin-left: 0.3rem; display: inline-block; transition: transform 0.2s; }
.toggle-arrow.open { transform: rotate(90deg); }
.details-row td { background: #fafafa; border-top: none !important; }
.sync-log { background: #1e1e1e; color: #d4d4d4; border-radius: 6px; padding: 1rem;
             font-family: 'Consolas', 'Courier New', monospace; font-size: 0.82rem;
             max-height: 400px; overflow-y: auto; line-height: 1.6; }
.sync-log .msg { margin: 0.1rem 0; }
.sync-log .msg::before { content: "> "; color: #008542; }
.sync-log .error { color: #f44; }
@media (max-width: 600px) {
  .search-box { flex-direction: column; }
  td, th { padding: 0.4rem; font-size: 0.8rem; }
}
"""

EMAIL_JS = """
function toggleEmails(btn, netid, firstName, lastName) {
    const panel = document.getElementById('emails-' + netid);
    if (panel.classList.contains('open')) {
        panel.classList.remove('open');
        btn.textContent = 'Show Emails';
        return;
    }
    panel.classList.add('open');
    btn.textContent = 'Hide Emails';
    const container = panel.querySelector('.email-container');
    if (container.dataset.loaded) return;

    container.innerHTML = '<div class="email-loading">Searching Outlook for emails from ' + firstName + ' ' + lastName + '...</div>';

    fetch('/api/emails?first=' + encodeURIComponent(firstName) + '&last=' + encodeURIComponent(lastName))
        .then(r => r.json())
        .then(data => {
            container.dataset.loaded = '1';
            if (data.error) {
                container.innerHTML = '<div class="email-none">Error: ' + data.error + '</div>';
                return;
            }
            if (data.emails.length === 0) {
                container.innerHTML = '<div class="email-none">No emails found from this student.</div>';
                return;
            }
            let html = '<div class="email-header"><h3>Emails from ' + firstName + ' ' + lastName + '</h3>';
            html += '<span class="email-count">' + data.emails.length + ' email' + (data.emails.length !== 1 ? 's' : '') + '</span></div>';
            html += '<ul class="email-list">';
            data.emails.forEach(e => {
                html += '<li class="email-item">';
                html += '<div><span class="email-subject">' + escapeHtml(e.subject) + '</span>';
                html += '<span class="email-date">' + e.date + '</span></div>';
                if (e.preview) {
                    html += '<div class="email-preview">' + escapeHtml(e.preview) + '</div>';
                }
                html += '</li>';
            });
            html += '</ul>';
            container.innerHTML = html;
        })
        .catch(err => {
            container.innerHTML = '<div class="email-none">Could not load emails. Is Outlook running?</div>';
        });
}

function toggleGrades(btn, netid, course, firstName, lastName) {
    const panel = document.getElementById('grades-' + netid + '-' + course);
    if (panel.classList.contains('open')) {
        panel.classList.remove('open');
        btn.textContent = 'Grades';
        return;
    }
    panel.classList.add('open');
    btn.textContent = 'Hide Grades';
    const container = panel.querySelector('.grades-container');
    if (container.dataset.loaded) return;

    container.innerHTML = '<div class="grades-loading">Loading assignment grades...</div>';

    fetch('/api/assignments?netid=' + encodeURIComponent(netid) + '&course=' + encodeURIComponent(course))
        .then(r => r.json())
        .then(data => {
            container.dataset.loaded = '1';
            if (data.error) {
                container.innerHTML = '<div class="grades-none">Error: ' + data.error + '</div>';
                return;
            }
            if (!data.assignments || data.assignments.length === 0) {
                container.innerHTML = '<div class="grades-none">No assignment details available.</div>';
                return;
            }
            let html = '<div class="grades-header"><h3>Assignments for ' + escapeHtml(firstName) + ' ' + escapeHtml(lastName) + '</h3></div>';
            html += '<table class="assign-table"><tr><th>Assignment</th><th>Score</th><th>Max</th><th>%</th></tr>';
            let totalScore = 0, totalMax = 0;
            data.assignments.forEach(a => {
                let scoreDisplay, pctDisplay, cls;
                if (a.score === '' || a.score === null) {
                    scoreDisplay = '<span class="score-none">&mdash;</span>';
                    pctDisplay = '<span class="score-none">&mdash;</span>';
                    cls = '';
                } else {
                    let s = parseFloat(a.score), m = parseFloat(a.max);
                    scoreDisplay = a.score;
                    totalScore += s;
                    totalMax += m;
                    if (m > 0) {
                        let pct = (s / m * 100);
                        pctDisplay = pct.toFixed(0) + '%';
                        cls = pct >= 80 ? 'score-high' : pct >= 60 ? 'score-mid' : 'score-low';
                    } else {
                        pctDisplay = '&mdash;';
                        cls = '';
                    }
                }
                html += '<tr><td>' + escapeHtml(a.name) + '</td>';
                html += '<td class="' + cls + '">' + scoreDisplay + '</td>';
                html += '<td>' + a.max + '</td>';
                html += '<td class="' + cls + '">' + pctDisplay + '</td></tr>';
            });
            // Total row
            let totalPct = totalMax > 0 ? (totalScore / totalMax * 100).toFixed(0) + '%' : '&mdash;';
            let totalCls = totalMax > 0 ? ((totalScore/totalMax*100) >= 80 ? 'score-high' : (totalScore/totalMax*100) >= 60 ? 'score-mid' : 'score-low') : '';
            html += '<tr><td><strong>Total</strong></td><td class="' + totalCls + '"><strong>' + totalScore.toFixed(1) + '</strong></td>';
            html += '<td><strong>' + totalMax.toFixed(1) + '</strong></td>';
            html += '<td class="' + totalCls + '"><strong>' + totalPct + '</strong></td></tr>';
            html += '</table>';
            container.innerHTML = html;
        })
        .catch(err => {
            container.innerHTML = '<div class="grades-none">Could not load assignments.</div>';
        });
}

function escapeHtml(str) {
    const div = document.createElement('div');
    div.textContent = str;
    return div.innerHTML;
}

function syncCourse(btn, courseKey) {
    if (btn.disabled) return;
    btn.disabled = true;
    btn.textContent = 'Syncing...';
    btn.style.opacity = '0.6';

    fetch('/api/sync-course', {
        method: 'POST',
        headers: {'Content-Type': 'application/json'},
        body: JSON.stringify({course: courseKey})
    })
    .then(r => r.json())
    .then(data => {
        if (data.error) {
            alert('Sync error: ' + data.error);
            btn.disabled = false;
            btn.textContent = 'Sync';
            btn.style.opacity = '1';
            return;
        }
        // Poll sync status
        const poll = setInterval(() => {
            fetch('/api/sync-status')
            .then(r => r.json())
            .then(st => {
                if (st.done) {
                    clearInterval(poll);
                    btn.disabled = false;
                    btn.style.opacity = '1';
                    if (st.error) {
                        btn.textContent = 'Failed';
                        btn.style.background = '#d32f2f';
                        alert('Sync failed: ' + st.error);
                    } else {
                        btn.textContent = 'Synced!';
                        btn.style.background = '#008542';
                        // Reload page after 1.5s to show updated grades
                        setTimeout(() => location.reload(), 1500);
                    }
                }
            });
        }, 3000);
    })
    .catch(err => {
        btn.disabled = false;
        btn.textContent = 'Sync';
        btn.style.opacity = '1';
        alert('Sync error: ' + err);
    });
}
"""

UPLOAD_JS = """
const zone = document.getElementById('dropZone');
const input = document.getElementById('fileInput');
const btn = document.getElementById('uploadBtn');
const list = document.getElementById('fileList');

zone.addEventListener('click', () => input.click());
zone.addEventListener('dragover', e => { e.preventDefault(); zone.classList.add('dragover'); });
zone.addEventListener('dragleave', () => zone.classList.remove('dragover'));
zone.addEventListener('drop', e => {
    e.preventDefault(); zone.classList.remove('dragover');
    input.files = e.dataTransfer.files; showFiles();
});
input.addEventListener('change', showFiles);

function showFiles() {
    const files = input.files;
    if (files.length) {
        list.innerHTML = Array.from(files).map(f =>
            '<div style="padding:0.3rem 0;color:#333;">' + f.name + '</div>'
        ).join('');
        btn.style.display = 'inline-block';
    }
}
"""

RULES_JS = """
let rulesPolling = null;

function setupRules() {
    const btn = document.getElementById('rulesBtn');
    const log = document.getElementById('rulesLog');
    btn.disabled = true;
    btn.textContent = 'Setting up...';
    log.style.display = 'block';
    log.innerHTML = '<div class="msg">Starting rule setup...</div>';

    fetch('/api/setup-rules', {method: 'POST'})
        .then(r => r.json())
        .then(data => {
            if (data.error) {
                addRulesLog(data.error, true);
                btn.disabled = false;
                btn.textContent = 'Sort Student Emails';
                return;
            }
            rulesPolling = setInterval(pollRulesStatus, 1000);
        });
}

function pollRulesStatus() {
    fetch('/api/rules-status')
        .then(r => r.json())
        .then(data => {
            const log = document.getElementById('rulesLog');
            log.innerHTML = data.messages.map(m =>
                '<div class="msg">' + rulesEscape(m) + '</div>'
            ).join('');
            log.scrollTop = log.scrollHeight;

            if (data.error) {
                addRulesLog('ERROR: ' + data.error, true);
            }
            if (data.done) {
                clearInterval(rulesPolling);
                const btn = document.getElementById('rulesBtn');
                btn.disabled = false;
                btn.textContent = 'Sort Student Emails';
            }
        });
}

function addRulesLog(msg, isError) {
    const log = document.getElementById('rulesLog');
    const cls = isError ? 'msg error' : 'msg';
    log.innerHTML += '<div class="' + cls + '">' + rulesEscape(msg) + '</div>';
    log.scrollTop = log.scrollHeight;
}

function rulesEscape(str) {
    const div = document.createElement('div');
    div.textContent = str;
    return div.innerHTML;
}
"""

SYNC_JS = """
let polling = null;

function startSync() {
    document.getElementById('syncBtn').disabled = true;
    document.getElementById('syncBtn').textContent = 'Syncing...';
    document.getElementById('syncLog').innerHTML = '<div class="msg">Starting grade sync...</div>';

    fetch('/api/sync-start', {method: 'POST'})
        .then(r => r.json())
        .then(data => {
            if (data.error) {
                addLog(data.error, true);
                return;
            }
            polling = setInterval(pollStatus, 1500);
        });
}

function pollStatus() {
    fetch('/api/sync-status')
        .then(r => r.json())
        .then(data => {
            const log = document.getElementById('syncLog');
            log.innerHTML = data.messages.map(m =>
                '<div class="msg">' + escapeHtml(m) + '</div>'
            ).join('');
            log.scrollTop = log.scrollHeight;

            if (data.error) {
                addLog('ERROR: ' + data.error, true);
            }
            if (data.done) {
                clearInterval(polling);
                document.getElementById('syncBtn').disabled = false;
                document.getElementById('syncBtn').textContent = 'Sync Grades';
                if (!data.error) addLog('Done!');
            }
        });
}

function addLog(msg, isError) {
    const log = document.getElementById('syncLog');
    const cls = isError ? 'msg error' : 'msg';
    log.innerHTML += '<div class="' + cls + '">' + escapeHtml(msg) + '</div>';
    log.scrollTop = log.scrollHeight;
}

function escapeHtml(str) {
    const div = document.createElement('div');
    div.textContent = str;
    return div.innerHTML;
}
"""


def grade_class(grade):
    """Return CSS class for grade coloring."""
    if not grade:
        return "grade-default"
    g = grade.strip().upper()
    if g.startswith("A"):
        return "grade-A"
    if g.startswith("B"):
        return "grade-B"
    if g.startswith("C"):
        return "grade-C"
    if g.startswith("D"):
        return "grade-D"
    if g.startswith("F"):
        return "grade-F"
    return "grade-default"


def base_html(title, body, active="", extra_js=""):
    search_cls = "active" if active == "search" else ""
    upload_cls = "active" if active == "upload" else ""
    sync_cls = "active" if active == "sync" else ""
    grading_cls = "active" if active == "grading" else ""
    flashes = get_flashed_messages(with_categories=True)
    flash_html = "".join(
        f'<div class="flash flash-{cat}">{msg}</div>' for cat, msg in flashes
    )
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title} – Student Lookup</title>
<style>{CSS}</style>
</head>
<body>
<nav class="navbar">
  <a href="/search" class="brand">Student Lookup</a>
  <a href="/search" class="nav-link {search_cls}">Search</a>
  <a href="/upload" class="nav-link {upload_cls}">Upload</a>
  <a href="/sync" class="nav-link {sync_cls}">Sync Grades</a>
  <a href="/grading" class="nav-link {grading_cls}">Grading</a>
  <span style="margin-left:auto;">
    <a href="/course/buan4320.502" class="nav-link">4320.502</a>
    <a href="/course/itss4351.003" class="nav-link">4351.003</a>
    <a href="/course/itss4351.002" class="nav-link">4351.002</a>
    <a href="/course/buan4320.501" class="nav-link">4320.501</a>
    <a href="/course/buan6320.s01" class="nav-link">6320.S01</a>
  </span>
</nav>
<div class="container">
  {flash_html}
  {body}
</div>
{f"<script>{extra_js}</script>" if extra_js else ""}
</body>
</html>"""

# ── Routes ───────────────────────────────────────────────────────────────────

def get_needs_grading():
    """Count 'Needs Grading' entries per course from assignments JSON in DB."""
    conn = get_db()
    rows = conn.execute(
        "SELECT course, assignments FROM students WHERE assignments LIKE '%Needs Grading%'"
    ).fetchall()
    conn.close()
    from collections import defaultdict
    course_ng = defaultdict(lambda: defaultdict(int))
    for r in rows:
        try:
            asns = json.loads(r['assignments'])
        except (json.JSONDecodeError, TypeError):
            continue
        for a in asns:
            if a.get('score', '').strip().lower() == 'needs grading':
                course_ng[r['course']][a['name']] = course_ng[r['course']][a['name']] + 1
    return course_ng


@app.route("/")
def index():
    return redirect(url_for("search"))

@app.route("/search")
def search():
    q = request.args.get("q", "").strip()
    results = []
    if q:
        conn = get_db()
        results = conn.execute("""
            SELECT last_name, first_name, netid, course, grade
            FROM students
            WHERE last_name LIKE ? OR first_name LIKE ? OR netid LIKE ?
            ORDER BY last_name, first_name, course
        """, (f"{q}%", f"{q}%", f"{q}%")).fetchall()
        conn.close()

    conn = get_db()
    total = conn.execute("SELECT COUNT(*) FROM students").fetchone()[0]
    course_count = conn.execute("SELECT COUNT(DISTINCT course) FROM students").fetchone()[0]
    conn.close()

    results_html = ""
    if q and results:
        seen_netids = set()
        rows = ""
        for r in results:
            netid = r['netid']
            fname = r['first_name']
            lname = r['last_name']
            grade = r['grade'] or ""
            email = f"{netid}@utdallas.edu"
            outlook = f"https://outlook.office365.com/mail/search/from:{email}"
            is_first = netid not in seen_netids
            seen_netids.add(netid)

            course_raw = r['course']
            email_btn = ""
            if is_first:
                email_btn = (
                    f'<button class="btn btn-sm btn-green" '
                    f"onclick=\"toggleEmails(this, '{netid}', '{fname}', '{lname}')\">"
                    f"Show Emails</button>"
                )

            grades_btn = (
                f'<button class="btn btn-sm btn-blue" '
                f"onclick=\"toggleGrades(this, '{netid}', '{course_raw}', '{fname}', '{lname}')\">"
                f"Grades</button>"
            )

            grade_html = ""
            if grade:
                grade_html = f'<span class="grade-tag {grade_class(grade)}">{grade}</span>'
            else:
                grade_html = '<span style="color:#ccc;">—</span>'

            bb_url = bb_gradebook_url(course_raw)
            bb_link = (
                f'<a href="{bb_url}" target="_blank" style="font-size:0.8rem; color:#c75b12;" '
                f"onclick=\"navigator.clipboard.writeText('{lname}')\">"
                f"BB Grades</a>"
            ) if bb_url else ""

            course_key_val = course_raw.lower().replace("-", ".").replace("_", ".").replace(" ", "")
            sync_btn = (
                f'<button class="btn btn-sm btn-sync" '
                f"onclick=\"syncCourse(this, '{course_key_val}')\">"
                f"Sync</button>"
            ) if bb_url else ""

            rows += f"""<tr>
                <td><strong>{lname}</strong></td>
                <td>{fname}</td>
                <td><code>{netid}</code></td>
                <td>{f'<a href="{bb_url}" target="_blank" class="course-tag" style="text-decoration:none;">{course_label(course_raw)}</a>' if bb_url else f'<span class="course-tag">{course_label(course_raw)}</span>'}</td>
                <td>{grade_html}</td>
                <td><a href="mailto:{email}">{email}</a></td>
                <td><a href="{outlook}" target="_blank" style="font-size:0.8rem;">Outlook</a> {bb_link}</td>
                <td>{grades_btn} {sync_btn} {email_btn}</td>
            </tr>
            <tr class="grades-panel" id="grades-{netid}-{course_raw}">
                <td colspan="8">
                    <div class="grades-container">
                        <div class="grades-loading">Loading assignments...</div>
                    </div>
                </td>
            </tr>
            <tr class="email-panel" id="emails-{netid}">
                <td colspan="8">
                    <div class="email-container">
                        <div class="email-loading">Loading emails from Outlook...</div>
                    </div>
                </td>
            </tr>"""

        count_label = f'{len(results)} result{"s" if len(results) != 1 else ""}'
        results_html = f"""
        <div class="card">
            <h2>{count_label} for "{q}"</h2>
            <table>
              <tr><th>Last Name</th><th>First Name</th><th>NetID</th><th>Course</th><th>Grade</th><th>Email</th><th></th><th></th></tr>
              {rows}
            </table>
        </div>"""

        return base_html("Search", f"""
        <div class="card">
          <h1>Search Students</h1>
          <form method="GET" action="/search" class="search-box">
            <input type="text" name="q" value="{q}" placeholder="Search by name or NetID..." autofocus>
            <button type="submit" class="btn">Search</button>
          </form>
          <p class="stats">{total} students across {course_count} course{"s" if course_count != 1 else ""}</p>
        </div>
        {results_html}
        """, active="search", extra_js=EMAIL_JS)
    elif q:
        results_html = '<div class="card"><p class="empty">No students found matching your search.</p></div>'

    return base_html("Search", f"""
    <div class="card">
      <h1>Search Students</h1>
      <form method="GET" action="/search" class="search-box">
        <input type="text" name="q" value="{q}" placeholder="Search by name or NetID..." autofocus>
        <button type="submit" class="btn">Search</button>
      </form>
      <p class="stats">{total} students across {course_count} course{"s" if course_count != 1 else ""}</p>
    </div>
    {results_html}
    """, active="search")

@app.route("/course/<ck>")
def course_dashboard(ck):
    """Show all students in a course with grades and needs-grading info."""
    conn = get_db()
    rows = conn.execute("""
        SELECT last_name, first_name, netid, course, grade, assignments
        FROM students
        WHERE LOWER(REPLACE(REPLACE(REPLACE(course, '-', '.'), '_', '.'), ' ', '')) = ?
        ORDER BY last_name, first_name
    """, (ck.lower(),)).fetchall()
    conn.close()

    if not rows:
        return base_html("Course Not Found", """
        <a href="/search" class="back-link">&larr; Back to Search</a>
        <div class="card"><p class="empty">No students found for this course.</p></div>
        """, active="search")

    course_raw = rows[0]['course']
    label = course_label(course_raw)
    ck_val = course_key(course_raw)
    bb_url = bb_gradebook_url(course_raw)
    total_students = len(rows)
    ng_students = 0
    table_rows = ""

    for r in rows:
        netid = r['netid']
        grade = r['grade'] or ""
        ng_items = []
        try:
            asns = json.loads(r['assignments']) if r['assignments'] else []
        except (json.JSONDecodeError, TypeError):
            asns = []
        for a in asns:
            if a.get('score', '').strip().lower() == 'needs grading':
                ng_items.append(a.get('name', 'Unknown'))

        if ng_items:
            ng_students += 1
            ng_cell = " ".join(f'<span class="ng-tag">{n}</span>' for n in ng_items)
        else:
            ng_cell = '<span class="ng-none">—</span>'

        grade_html = ""
        if grade:
            grade_html = f'<span class="grade-tag {grade_class(grade)}">{grade}</span>'
        else:
            grade_html = '<span style="color:#ccc;">—</span>'

        email = f"{netid}@utdallas.edu"
        table_rows += f"""<tr>
            <td><strong>{r['last_name']}</strong></td>
            <td>{r['first_name']}</td>
            <td><a href="/search?q={netid}"><code>{netid}</code></a></td>
            <td class="grade-col" style="display:none;">{grade_html}</td>
            <td>{ng_cell}</td>
            <td><a href="mailto:{email}">{email}</a></td>
        </tr>"""

    bb_link = ""
    if bb_url:
        bb_link = f'<a href="{bb_url}" target="_blank" style="font-size:0.85rem;">Open in Blackboard</a>'

    sync_btn = f"""<button class="btn btn-sm btn-sync" onclick="syncCourse(this, '{ck_val}')">Sync Grades</button>""" if bb_url else ""

    body = f"""
    <a href="/search" class="back-link">&larr; Back to Search</a>
    <div class="card">
        <div class="course-header">
            <h1 style="margin:0;">{label}</h1>
            <div style="display:flex; gap:0.5rem; align-items:center;">
                <button class="btn btn-sm" onclick="toggleGradeCol(this)">Show Grades</button>
                {sync_btn}
            </div>
        </div>
        <div class="course-summary">
            <span class="course-stat"><strong>{total_students}</strong> students</span>
            <span class="course-stat"><strong>{ng_students}</strong> with ungraded items</span>
            {f'<span class="course-stat">{bb_link}</span>' if bb_url else ''}
        </div>
        <table class="student-table">
            <tr>
                <th>Last Name</th><th>First Name</th><th>NetID</th>
                <th class="grade-col" style="display:none;">Grade</th><th>Needs Grading</th><th>Email</th>
            </tr>
            {table_rows}
        </table>
    </div>
    """

    course_js = """
function toggleGradeCol(btn) {
    const cols = document.querySelectorAll('.grade-col');
    const showing = cols[0] && cols[0].style.display !== 'none';
    cols.forEach(c => c.style.display = showing ? 'none' : '');
    btn.textContent = showing ? 'Show Grades' : 'Hide Grades';
}
""" + EMAIL_JS

    return base_html(label, body, active="search", extra_js=course_js)

@app.route("/api/emails")
def api_emails():
    first_name = request.args.get("first", "").strip()
    last_name = request.args.get("last", "").strip()
    if not last_name:
        return jsonify({"error": "Missing last name", "emails": []})
    try:
        emails = search_outlook_emails(first_name, last_name)
        return jsonify({"emails": emails})
    except Exception as e:
        return jsonify({"error": str(e), "emails": []})

@app.route("/api/assignments")
def api_assignments():
    netid = request.args.get("netid", "").strip()
    course = request.args.get("course", "").strip()
    if not netid or not course:
        return jsonify({"error": "Missing netid or course", "assignments": []})
    conn = get_db()
    row = conn.execute("SELECT assignments FROM students WHERE netid = ? AND course = ?",
                       (netid, course)).fetchone()
    conn.close()
    if not row or not row["assignments"]:
        return jsonify({"assignments": []})
    try:
        assignments = json.loads(row["assignments"])
        return jsonify({"assignments": assignments})
    except (json.JSONDecodeError, TypeError):
        return jsonify({"assignments": []})

@app.route("/sync")
def sync_page():
    return base_html("Sync Grades", """
    <div class="card">
      <h1>Sync Grades from Blackboard</h1>
      <p style="color:#666; margin-bottom:1rem;">
        Click the button below to open Chrome and sync grades from eLearning.
        The first time, you'll need to log in with your UTD credentials (SSO + Duo).
        After that, your session will be remembered.
      </p>
      <button class="btn btn-green" id="syncBtn" onclick="startSync()">Sync Grades</button>
      <div class="sync-log" id="syncLog" style="margin-top:1rem;">
        <div class="msg">Ready. Click "Sync Grades" to start.</div>
      </div>
    </div>
    <div class="card">
      <h2>Manual Upload</h2>
      <p style="color:#666; margin-bottom:0.5rem;">
        Alternatively, download the Full Grade Center from Blackboard (.xlsx) and upload it here.
        The file will be auto-detected and grades synced.
      </p>
      <a href="/upload" class="btn btn-sm">Go to Upload Page</a>
    </div>
    """, active="sync", extra_js=SYNC_JS)

@app.route("/api/sync-start", methods=["POST"])
def api_sync_start():
    if sync_status["running"]:
        return jsonify({"error": "Sync is already running"})
    t = threading.Thread(target=run_grade_sync, daemon=True)
    t.start()
    return jsonify({"ok": True})

@app.route("/api/sync-course", methods=["POST"])
def api_sync_course():
    """Sync grades for a single course by its course key (e.g. itss4351.003)."""
    if sync_status["running"]:
        return jsonify({"error": "A sync is already running"})
    course = request.json.get("course", "").strip().lower()
    bb_id = BB_COURSE_IDS.get(course)
    if not bb_id:
        return jsonify({"error": f"Unknown course: {course}"})
    # Find all course keys sharing this bb_id (merged courses)
    course_keys = [k for k, v in BB_COURSE_IDS.items() if v == bb_id]
    t = threading.Thread(target=run_single_course_sync, args=(bb_id, course_keys), daemon=True)
    t.start()
    return jsonify({"ok": True, "courses": course_keys})

@app.route("/api/sync-status")
def api_sync_status():
    return jsonify(sync_status)

@app.route("/api/setup-rules", methods=["POST"])
def api_setup_rules():
    if rules_status["running"]:
        return jsonify({"error": "Rule setup is already running"})
    t = threading.Thread(target=run_setup_outlook_rules, daemon=True)
    t.start()
    return jsonify({"ok": True})

@app.route("/api/rules-status")
def api_rules_status():
    return jsonify(rules_status)

@app.route("/upload", methods=["GET", "POST"])
def upload():
    if request.method == "POST":
        files = request.files.getlist("roster")
        if not files or all(f.filename == "" for f in files):
            flash("No files selected.", "error")
            return redirect(url_for("upload"))

        total_added = 0
        for f in files:
            if not f.filename.endswith((".xlsx", ".xls", ".csv")):
                flash(f"Skipped {f.filename} — not a supported file.", "error")
                continue
            try:
                file_data = f.read()
                class BytesFile:
                    def __init__(self, data, name):
                        self.data = data
                        self.filename = name
                    def read(self):
                        return self.data
                bf = BytesFile(file_data, f.filename)

                if f.filename.endswith(".csv") or _is_bb_tsv(file_data):
                    # Blackboard grade export (CSV or UTF-16 TSV .xls)
                    count, courses = load_master_grade(bf)
                    course_names = ", ".join(course_label(c) for c in sorted(courses))
                    flash(f"Loaded {count} students + grades from {f.filename} ({course_names})", "success")
                    total_added += count
                elif f.filename.endswith((".xlsx", ".xls")):
                    from io import BytesIO
                    if is_blackboard_grade_xlsx(file_data):
                        count = load_grade_xlsx(bf)
                        flash(f"Updated {count} grades from {f.filename}", "success")
                    else:
                        count = load_roster_from_file(bf)
                        total_added += count
                        flash(f"Loaded {count} students from {f.filename}", "success")
            except Exception as e:
                flash(f"Error processing {f.filename}: {e}", "error")

        if total_added:
            flash(f"Total: {total_added} students added/updated.", "success")
        return redirect(url_for("upload"))

    conn = get_db()
    courses = conn.execute("""
        SELECT course, COUNT(*) as cnt FROM students
        GROUP BY course ORDER BY course
    """).fetchall()
    conn.close()

    course_rows = ""
    for c in courses:
        course_rows += f"<tr><td><span class='course-tag'>{course_label(c['course'])}</span></td><td>{c['cnt']}</td></tr>"

    summary = ""
    if courses:
        summary = f"""
        <div class="card">
            <h2>Current Rosters</h2>
            <table>
              <tr><th>Course</th><th>Students</th></tr>
              {course_rows}
            </table>
        </div>"""

    rules_section = ""
    if courses:
        rules_section = """
        <div class="card">
            <h2>Auto-Sort Student Emails</h2>
            <p style="color:#666; margin-bottom:1rem;">
                Creates Outlook folders per course and moves student emails from your inbox.
                Also starts a background scanner that auto-sorts new emails every 5 minutes.
            </p>
            <button class="btn btn-green" id="rulesBtn" onclick="setupRules()">Sort Student Emails</button>
            <div class="sync-log" id="rulesLog" style="margin-top:1rem; display:none;"></div>
        </div>"""

    return base_html("Upload", f"""
    <div class="card">
      <h1>Upload Roster or Grades</h1>
      <p style="color:#666; margin-bottom:1rem;">
        Upload roster Excel files (.xlsx) or Blackboard grade exports (.xlsx/.csv).
        Grade files are auto-detected by their "Username" and "Total" columns.
      </p>
      <form method="POST" action="/upload" enctype="multipart/form-data" id="uploadForm">
        <div class="drop-zone" id="dropZone">
          <p style="font-size:1.5rem;">Drop files here</p>
          <p>Excel rosters (.xlsx) or Blackboard grade exports (.xlsx/.csv)</p>
          <input type="file" name="roster" multiple accept=".xlsx,.xls,.csv"
                 style="display:none" id="fileInput">
        </div>
        <div id="fileList" style="margin-top:1rem;"></div>
        <button type="submit" class="btn btn-green" style="margin-top:1rem; display:none;"
                id="uploadBtn">Upload &amp; Process</button>
      </form>
    </div>
    {summary}
    {rules_section}
    """, active="upload", extra_js=UPLOAD_JS + RULES_JS)

# ── Roster / Grade loading ───────────────────────────────────────────────────

def load_grade_csv(file_obj):
    """Load grades from a Blackboard CSV export, including all assignment columns."""
    import csv
    from io import StringIO

    data = file_obj.read()
    if isinstance(data, bytes):
        data = data.decode("utf-8-sig")
    reader = csv.reader(StringIO(data))
    headers = next(reader)

    # Find Username column
    username_idx = None
    for i, h in enumerate(headers):
        if h.strip('"').strip() == "Username":
            username_idx = i
            break
    if username_idx is None:
        raise ValueError("Could not find 'Username' column in CSV")

    # Find Total column
    total_idx = None
    for i, h in enumerate(headers):
        hl = h.lower()
        if "total" in hl and "score" in hl and "weighted" not in hl:
            total_idx = i
            break
    if total_idx is None:
        for i, h in enumerate(headers):
            if "Total" in h and "Weighted" not in h:
                total_idx = i
                break
    if total_idx is None:
        for i, h in enumerate(headers):
            if "Total" in h:
                total_idx = i
                break
    if total_idx is None:
        raise ValueError("Could not find 'Total' column in CSV")

    # Identify assignment columns
    assignment_cols = []
    for i, h in enumerate(headers):
        if i == total_idx:
            continue
        parsed = _parse_assignment_header(h)
        if parsed:
            assignment_cols.append((i, parsed[0], parsed[1]))

    conn = get_db()
    updated = 0
    for row in reader:
        if len(row) <= max(username_idx, total_idx):
            continue
        username = row[username_idx].strip()
        if not username:
            continue

        # Total grade
        score = row[total_idx].strip() if total_idx < len(row) else ""
        grade_str = ""
        if score:
            try:
                score_num = float(score)
                grade_str = f"{score_num:.1f}" if score_num != int(score_num) else str(int(score_num))
            except (ValueError, TypeError):
                grade_str = score

        # Assignments
        assignments = []
        for col_idx, name, max_pts in assignment_cols:
            if col_idx < len(row):
                cell_val = row[col_idx].strip()
                if cell_val:
                    try:
                        sv = float(cell_val)
                        score_str = f"{sv:.1f}" if sv != int(sv) else str(int(sv))
                    except (ValueError, TypeError):
                        score_str = cell_val
                else:
                    score_str = ""
            else:
                score_str = ""
            assignments.append({"name": name, "score": score_str, "max": max_pts})

        assignments_json = json.dumps(assignments) if assignments else ""
        conn.execute("UPDATE students SET grade = ?, assignments = ? WHERE netid = ?",
                     (grade_str, assignments_json, username))
        updated += 1

    conn.commit()
    conn.close()
    return updated


def is_blackboard_grade_xlsx(file_data):
    """Check if an xlsx file is a Blackboard grade export (vs a roster)."""
    import openpyxl
    from io import BytesIO

    wb = openpyxl.load_workbook(BytesIO(file_data), read_only=True)
    ws = wb.active
    # Check first row for Blackboard grade headers
    first_row = [str(cell.value or "").strip() for cell in list(ws.iter_rows(min_row=1, max_row=1))[0]]
    wb.close()
    has_username = "Username" in first_row
    has_total = any("Total" in h for h in first_row)
    return has_username and has_total


def _bb_course_id_to_key(bb_id):
    """Convert Blackboard course ID like '2262-UTDAL-BUAN-4351-SEC003-27401' to 'buan4351.003'."""
    import re
    m = re.search(r'(\w{4})-(\d{4})-SEC?(\w+)', bb_id, re.IGNORECASE)
    if m:
        dept = m.group(1).lower()
        num = m.group(2)
        sec = m.group(3).lower().lstrip("0") or "0"
        # Pad section: "1" -> "001", "s01" stays "s01"
        if sec.isdigit():
            sec = sec.zfill(3)
        return f"{dept}{num}.{sec}"
    return None


def _filename_to_course_key(filename):
    """Extract course key from filename like 'GRADE-BUAN-6320-SECS01.xls' or 'GRADE-BUAN-4320-SEC501.xls'."""
    import re
    base = os.path.splitext(os.path.basename(filename))[0].upper()
    # Match patterns like BUAN-6320-SECS01 or BUAN-4320-SEC501
    m = re.search(r'(\w{4})-(\d{4})-SEC(\w+)', base)
    if m:
        dept = m.group(1).lower()
        num = m.group(2)
        sec = m.group(3).lower()
        return f"{dept}{num}.{sec}"
    return None


def load_master_grade(file_obj):
    """Load a Blackboard grade export as master: creates students AND loads grades + assignments.

    Handles UTF-16 TSV (.xls), CSV, and detects merged courses via 'Child Course ID'.
    Course key is derived from Child Course ID column or filename.
    """
    import csv, re
    from io import StringIO

    data = file_obj.read()
    filename = getattr(file_obj, 'filename', '')

    # Detect format
    if data[:2] in (b'\xff\xfe', b'\xfe\xff'):
        text = data.decode('utf-16')
        delimiter = '\t'
    elif b'\t' in data[:500]:
        text = data.decode('utf-8-sig')
        delimiter = '\t'
    else:
        text = data.decode('utf-8-sig')
        delimiter = ','

    reader = csv.reader(StringIO(text), delimiter=delimiter)
    headers = [h.strip('"').strip() for h in next(reader)]

    # Find key columns
    def find_col(name):
        for i, h in enumerate(headers):
            if h == name:
                return i
        return None

    username_idx = find_col("Username")
    lastname_idx = find_col("Last Name")
    firstname_idx = find_col("First Name")
    child_course_idx = find_col("Child Course ID")

    if username_idx is None:
        raise ValueError("Could not find 'Username' column")
    if lastname_idx is None or firstname_idx is None:
        raise ValueError("Could not find name columns")

    # Find Total column
    total_idx = None
    for i, h in enumerate(headers):
        hl = h.lower()
        if "total" in hl and "score" in hl and "weighted" not in hl:
            total_idx = i
            break
    if total_idx is None:
        for i, h in enumerate(headers):
            if "Total" in h and "Weighted" not in h:
                total_idx = i
                break
    if total_idx is None:
        for i, h in enumerate(headers):
            if "Total" in h:
                total_idx = i
                break

    # Identify assignment columns
    assignment_cols = []
    for i, h in enumerate(headers):
        if i == total_idx:
            continue
        parsed = _parse_assignment_header(h)
        if parsed:
            assignment_cols.append((i, parsed[0], parsed[1]))

    # Fallback course key from filename
    fallback_course = _filename_to_course_key(filename)

    # Collect all rows first to determine courses for DELETE
    rows_data = []
    for row in reader:
        row = [cell.strip('"').strip() for cell in row]
        username = row[username_idx] if username_idx < len(row) else ""
        if not username:
            continue

        last_name = row[lastname_idx] if lastname_idx < len(row) else ""
        first_name = row[firstname_idx] if firstname_idx < len(row) else ""

        # Determine course key
        if child_course_idx is not None and child_course_idx < len(row):
            ck = _bb_course_id_to_key(row[child_course_idx])
        else:
            ck = None
        if not ck:
            ck = fallback_course
        if not ck:
            continue

        # Total grade
        grade_str = ""
        if total_idx is not None and total_idx < len(row):
            score = row[total_idx]
            if score:
                try:
                    sn = float(score)
                    grade_str = f"{sn:.1f}" if sn != int(sn) else str(int(sn))
                except (ValueError, TypeError):
                    grade_str = score

        # Assignments
        assignments = []
        for col_idx, name, max_pts in assignment_cols:
            cell_val = row[col_idx] if col_idx < len(row) else ""
            if cell_val:
                try:
                    sv = float(cell_val)
                    score_str = f"{sv:.1f}" if sv != int(sv) else str(int(sv))
                except (ValueError, TypeError):
                    score_str = cell_val
            else:
                score_str = ""
            assignments.append({"name": name, "score": score_str, "max": max_pts})

        rows_data.append((last_name, first_name, username, ck, grade_str,
                          json.dumps(assignments) if assignments else ""))

    # Delete old data for courses we're replacing, then insert fresh
    conn = get_db()
    courses_in_file = set(r[3] for r in rows_data)
    for ck in courses_in_file:
        conn.execute("DELETE FROM students WHERE course = ?", (ck,))

    added = 0
    for last_name, first_name, username, ck, grade_str, assignments_json in rows_data:
        conn.execute("""
            INSERT OR REPLACE INTO students (last_name, first_name, netid, course, grade, assignments)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (last_name, first_name, username, ck, grade_str, assignments_json))
        added += 1

    conn.commit()
    conn.close()
    return added, courses_in_file


def _is_bb_tsv(file_data):
    """Check if file is a Blackboard tab-separated UTF-16 export (.xls that's really TSV)."""
    try:
        if file_data[:2] in (b'\xff\xfe', b'\xfe\xff'):
            text = file_data.decode('utf-16')
        else:
            text = file_data.decode('utf-8-sig')
        first_line = text.split('\n')[0]
        return 'Username' in first_line and 'Total' in first_line and '\t' in first_line
    except Exception:
        return False


def load_grade_xls_tsv(file_obj):
    """Load grades from a Blackboard .xls export (tab-separated UTF-16).

    Same column layout as xlsx: Username, Total, and individual assignment columns.
    """
    import csv
    from io import StringIO

    data = file_obj.read()
    if data[:2] in (b'\xff\xfe', b'\xfe\xff'):
        text = data.decode('utf-16')
    else:
        text = data.decode('utf-8-sig')

    reader = csv.reader(StringIO(text), delimiter='\t')
    headers = next(reader)
    # Strip quotes that Blackboard wraps around headers
    headers = [h.strip('"').strip() for h in headers]

    # Find Username column
    username_idx = None
    for i, h in enumerate(headers):
        if h == "Username":
            username_idx = i
            break
    if username_idx is None:
        raise ValueError("Could not find 'Username' column")

    # Find Total column
    total_idx = None
    for i, h in enumerate(headers):
        hl = h.lower()
        if "total" in hl and "score" in hl and "weighted" not in hl:
            total_idx = i
            break
    if total_idx is None:
        for i, h in enumerate(headers):
            if "Total" in h and "Weighted" not in h:
                total_idx = i
                break
    if total_idx is None:
        for i, h in enumerate(headers):
            if "Total" in h:
                total_idx = i
                break
    if total_idx is None:
        raise ValueError("Could not find 'Total' column")

    # Identify assignment columns
    assignment_cols = []
    for i, h in enumerate(headers):
        if i == total_idx:
            continue
        parsed = _parse_assignment_header(h)
        if parsed:
            assignment_cols.append((i, parsed[0], parsed[1]))

    conn = get_db()
    updated = 0
    for row in reader:
        if len(row) <= max(username_idx, total_idx):
            continue
        # Strip quotes
        row = [cell.strip('"').strip() for cell in row]
        username = row[username_idx]
        if not username:
            continue

        # Total grade
        score = row[total_idx] if total_idx < len(row) else ""
        grade_str = ""
        if score:
            try:
                score_num = float(score)
                grade_str = f"{score_num:.1f}" if score_num != int(score_num) else str(int(score_num))
            except (ValueError, TypeError):
                grade_str = score.strip()

        # Assignments
        assignments = []
        for col_idx, name, max_pts in assignment_cols:
            if col_idx < len(row):
                cell_val = row[col_idx]
                if cell_val:
                    try:
                        sv = float(cell_val)
                        score_str = f"{sv:.1f}" if sv != int(sv) else str(int(sv))
                    except (ValueError, TypeError):
                        score_str = cell_val.strip()
                else:
                    score_str = ""
            else:
                score_str = ""
            assignments.append({"name": name, "score": score_str, "max": max_pts})

        assignments_json = json.dumps(assignments) if assignments else ""
        conn.execute("UPDATE students SET grade = ?, assignments = ? WHERE netid = ?",
                     (grade_str, assignments_json, username))
        updated += 1

    conn.commit()
    conn.close()
    return updated


def _parse_assignment_header(header):
    """Parse a Blackboard assignment header like 'SQL Install [Total Pts: 5 Score] |2522362'.
    Returns (clean_name, max_points) or None if not an assignment column."""
    import re
    m = re.match(r'^(.+?)\s*\[Total Pts:\s*([\d.]+)\s+Score\]\s*\|', header)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    # Also match headers with just [Total Pts: X] without the pipe/ID
    m2 = re.match(r'^(.+?)\s*\[Total Pts:\s*([\d.]+)', header)
    if m2:
        return m2.group(1).strip(), m2.group(2).strip()
    return None


def load_grade_xlsx(file_obj):
    """Load grades from a Blackboard xlsx grade export.

    Expected format: Row 1 = headers including 'Username' (netids) and
    a 'Total [Total Pts: up to X Score] |...' column with numeric scores.
    Also parses individual assignment columns (8+) into JSON.
    """
    import openpyxl
    from io import BytesIO

    data = file_obj.read()
    wb = openpyxl.load_workbook(BytesIO(data), read_only=True)
    ws = wb.active

    # Read headers from row 1
    headers = [str(cell.value or "").strip() for cell in list(ws.iter_rows(min_row=1, max_row=1))[0]]

    # Find Username column
    username_idx = None
    for i, h in enumerate(headers):
        if h == "Username":
            username_idx = i
            break
    if username_idx is None:
        wb.close()
        raise ValueError("Could not find 'Username' column in grade file")

    # Find Total column (look for "Total" with "Score" — the cumulative score, not "Weighted Total")
    total_idx = None
    for i, h in enumerate(headers):
        hl = h.lower()
        if "total" in hl and "score" in hl and "weighted" not in hl:
            total_idx = i
            break
    # Fallback: any column with "Total" in name
    if total_idx is None:
        for i, h in enumerate(headers):
            if "Total" in h and "Weighted" not in h:
                total_idx = i
                break
    # Last fallback: Weighted Total
    if total_idx is None:
        for i, h in enumerate(headers):
            if "Total" in h:
                total_idx = i
                break
    if total_idx is None:
        wb.close()
        raise ValueError("Could not find 'Total' column in grade file")

    # Identify individual assignment columns (typically columns 8+)
    # These have headers like "SQL Install [Total Pts: 5 Score] |2522362"
    assignment_cols = []  # list of (col_index, clean_name, max_points)
    for i, h in enumerate(headers):
        if i == total_idx:
            continue
        parsed = _parse_assignment_header(h)
        if parsed:
            assignment_cols.append((i, parsed[0], parsed[1]))

    conn = get_db()
    updated = 0
    for row in ws.iter_rows(min_row=2, values_only=True):
        if len(row) <= max(username_idx, total_idx):
            continue
        username = str(row[username_idx] or "").strip()
        score = row[total_idx]
        if not username:
            continue

        # Format total score
        grade_str = ""
        if score is not None and str(score).strip() != "":
            try:
                score_num = float(score)
                grade_str = f"{score_num:.1f}" if score_num != int(score_num) else str(int(score_num))
            except (ValueError, TypeError):
                grade_str = str(score).strip()

        # Build assignments JSON
        assignments = []
        for col_idx, name, max_pts in assignment_cols:
            if col_idx < len(row):
                cell_val = row[col_idx]
                if cell_val is not None and str(cell_val).strip() != "":
                    try:
                        sv = float(cell_val)
                        score_str = f"{sv:.1f}" if sv != int(sv) else str(int(sv))
                    except (ValueError, TypeError):
                        score_str = str(cell_val).strip()
                else:
                    score_str = ""
            else:
                score_str = ""
            assignments.append({"name": name, "score": score_str, "max": max_pts})

        assignments_json = json.dumps(assignments) if assignments else ""

        conn.execute("UPDATE students SET grade = ?, assignments = ? WHERE netid = ?",
                     (grade_str, assignments_json, username))
        updated += 1

    conn.commit()
    conn.close()
    wb.close()
    return updated

def load_roster_from_file(file_obj):
    import openpyxl
    from io import BytesIO

    data = file_obj.read()
    wb = openpyxl.load_workbook(BytesIO(data))
    ws = wb.active

    header_row = None
    for i, row in enumerate(ws.iter_rows(min_row=1, max_row=10, values_only=False), 1):
        vals = [cell.value for cell in row]
        if "NetId" in vals or "NetID" in vals or "netid" in [str(v).lower() for v in vals if v]:
            header_row = i
            break
    if header_row is None:
        raise ValueError("Could not find header row with 'NetId' column")

    headers = [cell.value for cell in ws[header_row]]
    col = {h: idx for idx, h in enumerate(headers) if h}

    conn = get_db()
    course_val = None
    count = 0

    for row in ws.iter_rows(min_row=header_row + 1, values_only=True):
        netid = row[col.get("NetId", col.get("NetID", 0))]
        if not netid:
            continue
        first_name = row[col.get("First_Name", 2)] or ""
        last_name = row[col.get("Last_Name", 4)] or ""
        class_val = row[col.get("Class", 6)] or ""
        course_k = course_key(class_val)

        if course_val is None:
            course_val = course_k
            conn.execute("DELETE FROM students WHERE course = ?", (course_k,))

        conn.execute("""
            INSERT OR REPLACE INTO students (last_name, first_name, netid, course)
            VALUES (?, ?, ?, ?)
        """, (last_name.strip(), first_name.strip(), str(netid).strip(), course_k))
        count += 1

    conn.commit()
    conn.close()
    return count

def load_roster_from_path(filepath):
    class FakeFile:
        def __init__(self, path):
            self.path = path
            self.filename = os.path.basename(path)
        def read(self):
            with open(self.path, "rb") as f:
                return f.read()
    return load_roster_from_file(FakeFile(filepath))

# ── Grading ──────────────────────────────────────────────────────────────────

GRADING_DIR = os.path.expanduser("~/Downloads")

# Course → assignment config: folder name, point total, assignment column header ID
GRADING_COURSES = {
    "buan6320": {
        "label": "BUAN 6320",
        "css": "btn-buan6320",
        "assignments": {
            "assignment2": {
                "label": "Assignment 2 - SQL DML",
                "folder": "6320",
                "total_pts": 20,
                "bb_col": "Assignment 2 - SQL DML [Total Pts: 20 Score] |2522365",
            },
        },
    },
    "buan4320-s01": {
        "label": "BUAN 4320.501",
        "css": "btn-buan4320-s01",
        "assignments": {
            "assignment2": {
                "label": "Assignment 2 - SQL DML",
                "folder": "4320-S01",
                "total_pts": 30,
                "bb_col": "",
            },
        },
    },
    "buan4320-s02": {
        "label": "BUAN 4320.502",
        "css": "btn-buan4320-s02",
        "assignments": {
            "assignment2": {
                "label": "Assignment 2 - SQL DML",
                "folder": "4320-S02",
                "total_pts": 30,
                "bb_col": "",
            },
        },
    },
}

# In-memory grading results cache
grading_results = {}
grading_status = {"running": False, "messages": [], "done": False, "error": None}


def grading_log(msg):
    grading_status["messages"].append(msg)
    print(f"  [grading] {msg}")


def grade_sql_assignment(folder_path, total_pts, skip_ids=None):
    """Grade all .docx files in folder_path. Returns list of {student_id, score, deductions, comment}.
    skip_ids: set of uppercase student IDs to skip (already in Excel)."""
    import re
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn as docx_qn

    if skip_ids is None:
        skip_ids = set()

    results = []
    docx_files = glob_mod.glob(os.path.join(folder_path, "*.docx"))

    if not docx_files:
        return results

    for filepath in docx_files:
        filename = os.path.basename(filepath)
        # Skip already-graded files
        if "GRADED" in filename.upper():
            continue
        # Extract student ID from filename (prefix before _ or -)
        m = re.match(r'([A-Za-z]+\d+)', filename)
        student_id = m.group(1).upper() if m else filename.replace('.docx', '')

        # Skip if already in Excel
        if student_id in skip_ids:
            grading_log(f"Skipping {filename} ({student_id}) — already graded in Excel")
            continue

        grading_log(f"Grading {filename} ({student_id})...")

        try:
            doc = DocxDocument(filepath)
            paras = [p.text.strip() for p in doc.paragraphs]

            # Count images
            img_count = 0
            for p in doc.paragraphs:
                imgs = p._element.findall('.//' + docx_qn('wp:inline'))
                img_count += len(imgs)

            # Find question paragraphs and their SQL queries
            questions = {}  # q_num -> {query_lines, has_screenshot, join_type, ...}
            current_q = None
            in_query = False
            query_lines = []

            for i, text in enumerate(paras):
                # Detect question headers
                qm = re.match(r'Question\s+(\d+)', text)
                if qm:
                    if current_q and query_lines:
                        questions[current_q]["query"] = " ".join(query_lines)
                    current_q = int(qm.group(1))
                    questions[current_q] = {"query": "", "has_output": False}
                    in_query = False
                    query_lines = []
                    continue

                if text == "Your SQL Query:":
                    in_query = True
                    query_lines = []
                    continue

                if text.startswith("Your Output") or text.startswith("Verification"):
                    if current_q and query_lines:
                        questions[current_q]["query"] = " ".join(query_lines)
                    in_query = False
                    continue

                if in_query and text:
                    query_lines.append(text)

            # Last question
            if current_q and query_lines:
                questions[current_q]["query"] = " ".join(query_lines)

            # Check for screenshots near output sections
            for p_idx, p in enumerate(doc.paragraphs):
                if p.text.strip().startswith("Your Output"):
                    for j in range(p_idx, min(p_idx + 5, len(doc.paragraphs))):
                        imgs = doc.paragraphs[j]._element.findall('.//' + docx_qn('wp:inline'))
                        if imgs:
                            # Find which question this belongs to
                            for qi in range(p_idx, -1, -1):
                                qm2 = re.match(r'Question\s+(\d+)', doc.paragraphs[qi].text.strip())
                                if qm2:
                                    qn = int(qm2.group(1))
                                    if qn in questions:
                                        questions[qn]["has_output"] = True
                                    break
                            break

            # Check verification outputs
            verif_present = set()
            for p_idx, p in enumerate(doc.paragraphs):
                if "Verification" in p.text and "Requirement" not in p.text and "query:" not in p.text:
                    has_img = False
                    for j in range(p_idx, min(p_idx + 4, len(doc.paragraphs))):
                        imgs = doc.paragraphs[j]._element.findall('.//' + docx_qn('wp:inline'))
                        if imgs:
                            has_img = True
                            break
                    if has_img:
                        # Find which question this verification belongs to
                        for qi in range(p_idx, -1, -1):
                            qm3 = re.match(r'Question\s+(\d+)', doc.paragraphs[qi].text.strip())
                            if qm3:
                                verif_present.add(int(qm3.group(1)))
                                break

            # Grade each question
            deductions = []
            score = float(total_pts)
            pts_per_q = total_pts / 20.0  # 1 point per question for 20 questions

            for qn_num in range(1, 21):
                q = questions.get(qn_num)
                if not q:
                    deductions.append(f"Q{qn_num}: -{pts_per_q} (missing)")
                    score -= pts_per_q
                    continue

                query_upper = q["query"].upper()

                # Check join type issues
                # Q7: Stock Clerk - should use INNER JOIN not LEFT JOIN
                if qn_num == 7 and "LEFT JOIN" in query_upper and "STOCK CLERK" in query_upper:
                    deductions.append(f"Q{qn_num}: -0.5 (LEFT JOIN instead of INNER JOIN)")
                    score -= 0.5

                # Q11: Hire date filter - should use INNER JOIN not LEFT JOIN
                if qn_num == 11 and "LEFT JOIN" in query_upper and "1994" in q["query"]:
                    deductions.append(f"Q{qn_num}: -0.5 (LEFT JOIN instead of INNER JOIN)")
                    score -= 0.5

                # Q5: Countries in Americas - should use INNER JOIN not LEFT JOIN
                if qn_num == 5 and "LEFT JOIN" in query_upper and "AMERICA" in query_upper:
                    deductions.append(f"Q{qn_num}: -0.5 (LEFT JOIN instead of INNER JOIN)")
                    score -= 0.5

                # Q1: Cross Join should not have ON clause
                if qn_num == 1 and "CROSS JOIN" in query_upper and " ON " in query_upper:
                    deductions.append(f"Q{qn_num}: -0.5 (ON clause in CROSS JOIN)")
                    score -= 0.5

                # Q1: Cross Join should not have LIMIT
                if qn_num == 1 and "CROSS JOIN" in query_upper and "LIMIT" in query_upper:
                    deductions.append(f"Q{qn_num}: -0.5 (LIMIT restricts CROSS JOIN results)")
                    score -= 0.5

                # Q8: Self-join for employee/manager - should use INNER JOIN not LEFT JOIN
                if qn_num == 8 and "LEFT JOIN" in query_upper and "MANAGER" in query_upper:
                    deductions.append(f"Q{qn_num}: -0.5 (LEFT JOIN instead of INNER JOIN for self-join)")
                    score -= 0.5

                # Q14: Full outer join simulation check
                if qn_num == 14:
                    if "UNION" in query_upper:
                        # Check for incorrect WHERE IS NULL pattern
                        if "WHERE" in query_upper and "IS NULL" in query_upper and "UNION" in query_upper:
                            # If both sides have WHERE IS NULL, they only return unmatched rows
                            parts = query_upper.split("UNION")
                            if len(parts) >= 2:
                                left_has_null = "IS NULL" in parts[0]
                                right_has_null = "IS NULL" in parts[1]
                                if left_has_null and right_has_null:
                                    deductions.append(f"Q{qn_num}: -1.0 (incorrect FULL OUTER JOIN - only returns unmatched rows)")
                                    score -= 1.0
                    else:
                        deductions.append(f"Q{qn_num}: -1.0 (missing UNION for FULL OUTER JOIN simulation)")
                        score -= 1.0

                # Q15: Self-join should use INNER JOIN (Q16 is where LEFT JOIN is needed)
                if qn_num == 15 and "LEFT JOIN" in query_upper:
                    deductions.append(f"Q{qn_num}: -0.5 (LEFT JOIN instead of INNER JOIN - Q16 requires LEFT JOIN)")
                    score -= 0.5

                # Q16: Should have ORDER BY
                if qn_num == 16 and "ORDER BY" not in query_upper:
                    deductions.append(f"Q{qn_num}: -0.5 (missing ORDER BY last_name)")
                    score -= 0.5

            # Check missing verifications (required after Q3, Q7, Q12, Q16, Q19)
            for vq in [3, 7, 12, 16, 19]:
                if vq not in verif_present:
                    deductions.append(f"Q{vq}: -0.5 (missing verification output)")
                    score -= 0.5

            score = max(0, round(score * 2) / 2)  # round to nearest 0.5
            comment = " | ".join(deductions) if deductions else "Perfect score"

            # Build a set of deducted question numbers for marking
            deducted_qs = set()
            for d in deductions:
                dq = re.match(r'Q(\d+)', d)
                if dq:
                    deducted_qs.add(int(dq.group(1)))

            # Create graded docx with green ✓ / red marks
            graded_filename = filename.replace('.docx', '_GRADED.docx')
            graded_path = os.path.join(folder_path, graded_filename)
            try:
                from docx.shared import RGBColor, Pt
                graded_doc = DocxDocument(filepath)
                for p in graded_doc.paragraphs:
                    qm_g = re.match(r'(Question\s+(\d+))', p.text.strip())
                    if qm_g:
                        qn_g = int(qm_g.group(2))
                        if qn_g in deducted_qs:
                            # Find matching deduction text
                            ded_texts = [d for d in deductions if d.startswith(f"Q{qn_g}:")]
                            mark_text = "  " + "; ".join(ded_texts)
                            run = p.add_run(mark_text)
                            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                            run.font.size = Pt(10)
                            run.bold = True
                        else:
                            run = p.add_run("  ✓")
                            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                            run.font.size = Pt(12)
                            run.bold = True
                # Add score summary at end
                score_para = graded_doc.add_paragraph()
                sr = score_para.add_run(f"\nFinal Score: {score}/{total_pts}")
                sr.font.size = Pt(14)
                sr.bold = True
                sr.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
                if deductions:
                    dp = graded_doc.add_paragraph()
                    dr = dp.add_run("Deductions: " + " | ".join(deductions))
                    dr.font.size = Pt(10)
                    dr.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

                graded_doc.save(graded_path)
                grading_log(f"  Saved {graded_filename}")
            except Exception as ge:
                graded_path = ""
                grading_log(f"  Warning: could not create graded doc: {ge}")

            results.append({
                "student_id": student_id,
                "filename": filename,
                "score": score,
                "total": total_pts,
                "deductions": deductions,
                "comment": comment,
                "img_count": img_count,
                "q_count": len(questions),
                "verif_count": len(verif_present),
                "graded_file": graded_filename if graded_path else "",
                "graded_path": graded_path,
            })

            grading_log(f"  {student_id}: {score}/{total_pts}")

        except Exception as e:
            grading_log(f"  Error grading {filename}: {e}")
            results.append({
                "student_id": student_id,
                "filename": filename,
                "score": 0,
                "total": total_pts,
                "deductions": [f"Error: {str(e)}"],
                "comment": f"Error: {str(e)}",
                "img_count": 0,
                "q_count": 0,
            })

    return results


def run_grading(course_key, assign_key):
    """Background thread: grade assignments and store results."""
    global grading_status, grading_results
    grading_status = {"running": True, "messages": [], "done": False, "error": None}

    try:
        course = GRADING_COURSES.get(course_key)
        if not course:
            grading_status["error"] = f"Unknown course: {course_key}"
            return
        assign = course["assignments"].get(assign_key)
        if not assign:
            grading_status["error"] = f"Unknown assignment: {assign_key}"
            return

        folder = os.path.join(GRADING_DIR, assign["folder"])
        if not os.path.isdir(folder):
            grading_status["error"] = f"Folder not found: {folder}"
            return

        grading_log(f"Grading {course['label']} - {assign['label']}")
        grading_log(f"Folder: {folder}")

        # Load existing Excel to find already-graded student IDs
        excel_path = os.path.join(folder, f"{assign_key}_Grade_Upload.xlsx")
        already_graded = set()
        try:
            import openpyxl
            if os.path.exists(excel_path):
                ewb = openpyxl.load_workbook(excel_path)
                ews = ewb.active
                for row in ews.iter_rows(min_row=2, max_col=1, values_only=True):
                    if row[0]:
                        already_graded.add(str(row[0]).strip().upper())
                ewb.close()
                if already_graded:
                    grading_log(f"Found {len(already_graded)} already-graded students in Excel — skipping them")
        except Exception as e:
            grading_log(f"Note: could not read existing Excel: {e}")

        results = grade_sql_assignment(folder, assign["total_pts"], skip_ids=already_graded)
        cache_key = f"{course_key}_{assign_key}"
        grading_results[cache_key] = results

        # Append new grades to Excel (or create fresh if none exists)
        try:
            import openpyxl
            from openpyxl.styles import Font as XlFont
            if os.path.exists(excel_path) and already_graded:
                wb = openpyxl.load_workbook(excel_path)
                ws = wb.active
                next_row = ws.max_row + 1
            else:
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "Grades"
                ws["A1"] = "Username"
                ws["B1"] = assign["bb_col"]
                ws["C1"] = "Grading Notes"
                ws["D1"] = "Screenshots"
                ws["E1"] = "Verifications"
                for cell in [ws["A1"], ws["B1"], ws["C1"], ws["D1"], ws["E1"]]:
                    cell.font = XlFont(bold=True)
                next_row = 2
            for i, r in enumerate(results, next_row):
                ws.cell(row=i, column=1, value=r["student_id"].lower())
                ws.cell(row=i, column=2, value=r["score"])
                ws.cell(row=i, column=3, value=r["comment"])
                ws.cell(row=i, column=4, value=r.get("img_count", 0))
                ws.cell(row=i, column=5, value=r.get("verif_count", 0))
            ws.column_dimensions["A"].width = 15
            ws.column_dimensions["B"].width = 50
            ws.column_dimensions["C"].width = 100
            ws.column_dimensions["D"].width = 14
            ws.column_dimensions["E"].width = 14
            wb.save(excel_path)
            grading_log(f"Excel saved: {excel_path}")
        except Exception as e:
            grading_log(f"Excel export error: {e}")

        # Move GRADED files to a subfolder
        try:
            graded_dir = os.path.join(folder, "BUAN6320")
            os.makedirs(graded_dir, exist_ok=True)
            graded_files = glob_mod.glob(os.path.join(folder, "*_GRADED.docx"))
            moved = 0
            for gf in graded_files:
                dest = os.path.join(graded_dir, os.path.basename(gf))
                import shutil
                shutil.move(gf, dest)
                moved += 1
            if moved:
                grading_log(f"Moved {moved} graded files to {graded_dir}")
                # Update graded_path in results to point to new location
                for r in results:
                    if r.get("graded_file"):
                        r["graded_path"] = os.path.join(graded_dir, r["graded_file"])
        except Exception as e:
            grading_log(f"Warning: could not move graded files: {e}")

        grading_log(f"Done! {len(results)} new assignments graded.")

    except Exception as e:
        grading_status["error"] = str(e)
        grading_log(f"Error: {e}")
    finally:
        grading_status["running"] = False
        grading_status["done"] = True


GRADING_JS = """
function toggleStudentIds() {
    const cells = document.querySelectorAll('.sid-cell');
    const header = document.querySelector('.sid-header');
    const btnText = document.getElementById('sidBtnText');
    const showing = cells.length > 0 && cells[0].style.display !== 'none';
    cells.forEach(c => c.style.display = showing ? 'none' : 'table-cell');
    if (header) header.style.display = showing ? 'none' : 'table-cell';
    if (btnText) btnText.textContent = showing ? 'Show Student IDs' : 'Hide Student IDs';
}

function gradeAssignments(courseKey, assignKey) {
    const btn = document.getElementById('gradeBtn');
    const log = document.getElementById('gradingLog');
    const results = document.getElementById('gradingResults');
    btn.disabled = true;
    btn.textContent = 'Grading...';
    log.style.display = 'block';
    log.innerHTML = '<div class="msg">Starting grading...</div>';

    fetch('/api/grade-start?course=' + courseKey + '&assignment=' + assignKey, {method: 'POST'})
        .then(r => r.json())
        .then(data => {
            if (data.error) {
                log.innerHTML += '<div class="msg" style="color:#f44;">' + data.error + '</div>';
                btn.disabled = false;
                btn.textContent = 'Grade Assignments';
                return;
            }
            let poll = setInterval(() => {
                fetch('/api/grade-status')
                    .then(r => r.json())
                    .then(s => {
                        log.innerHTML = s.messages.map(m =>
                            '<div class="msg">' + escapeHtml(m) + '</div>'
                        ).join('');
                        log.scrollTop = log.scrollHeight;
                        if (s.done) {
                            clearInterval(poll);
                            btn.disabled = false;
                            btn.textContent = 'Re-Grade';
                            // Reload page to show results
                            setTimeout(() => location.reload(), 500);
                        }
                    });
            }, 1000);
        });
}

function escapeHtml(str) {
    const div = document.createElement('div');
    div.textContent = str;
    return div.innerHTML;
}

// Drag and drop for grading folder
const gZone = document.getElementById('gradingDrop');
if (gZone) {
    const gInput = document.getElementById('gradingFiles');
    gZone.addEventListener('click', () => gInput.click());
    gZone.addEventListener('dragover', e => { e.preventDefault(); gZone.classList.add('dragover'); });
    gZone.addEventListener('dragleave', () => gZone.classList.remove('dragover'));
    gZone.addEventListener('drop', e => {
        e.preventDefault(); gZone.classList.remove('dragover');
        gInput.files = e.dataTransfer.files;
        document.getElementById('gradingUploadForm').submit();
    });
    gInput.addEventListener('change', () => {
        document.getElementById('gradingUploadForm').submit();
    });
}
"""


@app.route("/grading")
def grading_page():
    # Show course buttons
    course_btns = ""
    for ck, cv in GRADING_COURSES.items():
        course_btns += f'<a href="/grading/{ck}" class="course-btn {cv["css"]}">{cv["label"]}</a> '

    return base_html("Grading", f"""
    <div class="card">
      <h1>Assignment Grading</h1>
      <p style="color:#666; margin-bottom:1rem;">Select a course to grade assignments.</p>
      <div>{course_btns}</div>
    </div>
    """, active="grading")


@app.route("/grading/<course_key>")
def grading_course(course_key):
    course = GRADING_COURSES.get(course_key)
    if not course:
        flash("Unknown course.", "error")
        return redirect(url_for("grading_page"))

    # Course buttons with active highlight
    course_btns = ""
    for ck, cv in GRADING_COURSES.items():
        active = " course-btn-active" if ck == course_key else ""
        course_btns += f'<a href="/grading/{ck}" class="course-btn {cv["css"]}{active}">{cv["label"]}</a> '

    # Assignment buttons
    assign_btns = ""
    for ak, av in course["assignments"].items():
        assign_btns += f'<a href="/grading/{course_key}/{ak}" class="assign-btn">{av["label"]}</a> '

    return base_html(f"Grading - {course['label']}", f"""
    <div class="card">
      <h1>Assignment Grading</h1>
      <div>{course_btns}</div>
    </div>
    <div class="card">
      <h2>{course['label']} - Select Assignment</h2>
      <div>{assign_btns}</div>
    </div>
    """, active="grading")


@app.route("/grading/<course_key>/<assign_key>")
def grading_assignment(course_key, assign_key):
    course = GRADING_COURSES.get(course_key)
    if not course:
        flash("Unknown course.", "error")
        return redirect(url_for("grading_page"))
    assign = course["assignments"].get(assign_key)
    if not assign:
        flash("Unknown assignment.", "error")
        return redirect(url_for("grading_course", course_key=course_key))

    # Course buttons
    course_btns = ""
    for ck, cv in GRADING_COURSES.items():
        active = " course-btn-active" if ck == course_key else ""
        course_btns += f'<a href="/grading/{ck}" class="course-btn {cv["css"]}{active}">{cv["label"]}</a> '

    # Assignment buttons
    assign_btns = ""
    for ak, av in course["assignments"].items():
        active = " assign-btn-active" if ak == assign_key else ""
        assign_btns += f'<a href="/grading/{course_key}/{ak}" class="assign-btn{active}">{av["label"]}</a> '

    # Check folder for ungraded files
    folder = os.path.join(GRADING_DIR, assign["folder"])
    docx_count = len([f for f in glob_mod.glob(os.path.join(folder, "*.docx")) if "GRADED" not in os.path.basename(f).upper()]) if os.path.isdir(folder) else 0

    # Load all graded results from Excel file (persistent source of truth)
    excel_path = os.path.join(folder, f"{assign_key}_Grade_Upload.xlsx")
    all_results = []
    if os.path.exists(excel_path):
        try:
            import openpyxl
            ewb = openpyxl.load_workbook(excel_path)
            ews = ewb.active
            for row in ews.iter_rows(min_row=2, values_only=True):
                sid = str(row[0]).strip().upper() if row[0] else ""
                score_val = float(row[1]) if row[1] is not None else 0
                comment = str(row[2]) if row[2] else ""
                try:
                    img_count = int(row[3]) if len(row) > 3 and row[3] is not None else 0
                except (ValueError, TypeError):
                    img_count = 0
                try:
                    verif_count = int(row[4]) if len(row) > 4 and row[4] is not None else 0
                except (ValueError, TypeError):
                    verif_count = 0
                if sid:
                    # Check if graded file exists in BUAN6320 subfolder
                    graded_file = ""
                    graded_subdir = os.path.join(folder, "BUAN6320")
                    if os.path.isdir(graded_subdir):
                        matches = glob_mod.glob(os.path.join(graded_subdir, f"{sid}*_GRADED.docx"))
                        if not matches:
                            matches = glob_mod.glob(os.path.join(graded_subdir, f"{sid.lower()}*_GRADED.docx"))
                        if matches:
                            graded_file = os.path.basename(matches[0])
                    # Fallback: find original submission file in main folder
                    orig_file = ""
                    if not graded_file:
                        sid_lower = sid.lower()
                        for fn in os.listdir(folder):
                            fn_lower = fn.lower()
                            if fn_lower.startswith(f"assignment 2") and sid_lower in fn_lower and not fn_lower.endswith(".txt"):
                                ext = os.path.splitext(fn)[1].lower()
                                if ext in (".docx", ".pdf", ".sql"):
                                    orig_file = fn
                                    if ext == ".docx":
                                        break  # prefer docx
                    all_results.append({
                        "student_id": sid,
                        "score": score_val,
                        "total": assign["total_pts"],
                        "comment": comment,
                        "img_count": img_count,
                        "verif_count": verif_count,
                        "graded_file": graded_file,
                        "orig_file": orig_file,
                    })
            ewb.close()
        except Exception as e:
            pass  # Excel not readable, show empty

    # Results table
    results_html = ""
    if all_results:
        rows = ""
        for r in sorted(all_results, key=lambda x: x["student_id"]):
            score = r["score"]
            total = r["total"]
            if score == total:
                score_cls = "score-perfect"
            elif score >= total * 0.9:
                score_cls = "score-good"
            else:
                score_cls = "score-deducted"

            comment = r["comment"]
            if comment and comment != "Perfect score":
                deductions = comment.split(" | ")
                deductions_html = "<br>".join(
                    f'<span style="color:#721c24;">{d}</span>' for d in deductions
                )
            else:
                deductions_html = '<span style="color:#155724;">No deductions</span>'

            graded_link = ""
            if r.get("graded_file"):
                graded_link = f'<a href="/api/graded-file?course={course_key}&assignment={assign_key}&file={r["graded_file"]}" target="_blank">View Graded</a>'
            elif r.get("orig_file"):
                orig_quoted = urllib_parse.quote(r["orig_file"])
                graded_link = f'<a href="/api/submission-file?course={course_key}&assignment={assign_key}&file={orig_quoted}" target="_blank" style="color:#e65100;">View Original</a>'

            rows += f"""<tr>
                <td class="sid-cell" style="display:none;"><code><strong>{r['student_id']}</strong></code></td>
                <td class="score-cell {score_cls}">{score}/{total}</td>
                <td class="comment-cell">{deductions_html}</td>
                <td style="text-align:center;">{r.get('img_count', 0)}</td>
                <td style="text-align:center;">{r.get('verif_count', 0)}</td>
                <td style="text-align:center;">{graded_link}</td>
            </tr>"""

        avg_score = sum(r["score"] for r in all_results) / len(all_results)
        results_html = f"""
        <div class="card grade-result">
            <h2>Grading Results ({len(all_results)} submissions)</h2>
            <p style="color:#666; margin-bottom:0.5rem;">
                Average: <strong>{avg_score:.1f}/{assign['total_pts']}</strong> |
                <a href="/api/grade-download?course={course_key}&assignment={assign_key}">Download Excel for Blackboard</a>
            </p>
            <table>
              <tr>
                <th class="sid-header" style="display:none; cursor:pointer;" onclick="toggleStudentIds()">Student ID <span id="sidArrow" class="toggle-arrow open">&#9654;</span></th>
                <th>Score</th><th>Deductions</th><th>Screenshots</th><th>Verifications</th><th>Graded File</th>
              </tr>
              {rows}
            </table>
            <button class="btn" style="margin-top:0.8rem; font-size:0.85rem;" onclick="toggleStudentIds()">
                <span id="sidBtnText">Show Student IDs</span>
            </button>
        </div>"""

    folder_info = f"{docx_count} new .docx files to grade" if docx_count > 0 else "No new files to grade"

    # Build file browser listing all student submission files
    files_html = ""
    if os.path.isdir(folder):
        submission_files = []
        for fn in sorted(os.listdir(folder)):
            if fn.startswith("Assignment 2") or fn.startswith("assignment2"):
                fp = os.path.join(folder, fn)
                if os.path.isfile(fp):
                    ext = os.path.splitext(fn)[1].lower()
                    # Extract student ID from filename
                    parts = fn.split("_")
                    sid = parts[1] if len(parts) > 1 else ""
                    size_kb = os.path.getsize(fp) / 1024
                    submission_files.append((fn, sid, ext, size_kb))

        # Also check graded subfolder
        graded_dir = os.path.join(folder, "assignment2_graded")
        graded_files = []
        if os.path.isdir(graded_dir):
            for fn in sorted(os.listdir(graded_dir)):
                fp = os.path.join(graded_dir, fn)
                if os.path.isfile(fp):
                    ext = os.path.splitext(fn)[1].lower()
                    size_kb = os.path.getsize(fp) / 1024
                    graded_files.append((fn, ext, size_kb))

        # Group submissions by student
        students = {}
        for fn, sid, ext, size_kb in submission_files:
            if ext == ".txt":
                continue  # skip metadata txt files
            if sid not in students:
                students[sid] = []
            students[sid].append((fn, ext, size_kb))

        if students:
            file_rows = ""
            for sid in sorted(students.keys()):
                file_links = ""
                for fn, ext, size_kb in students[sid]:
                    icon = {"docx": "📄", ".pdf": "📕", ".sql": "📝", ".doc": "📄"}.get(ext, "📎")
                    if ext == ".docx":
                        icon = "📄"
                    elif ext == ".pdf":
                        icon = "📕"
                    elif ext == ".sql":
                        icon = "📝"
                    else:
                        icon = "📎"
                    view_url = f"/api/submission-file?course={course_key}&assignment={assign_key}&file={urllib_parse.quote(fn)}"
                    # Extract just the original filename: everything after the last netid_ in the attempt portion
                    short_name = fn
                    if "_attempt_" in fn:
                        after_attempt = fn.split("_attempt_")[1]
                        # Format: date_netid_originalfile.ext — skip date and netid prefix
                        attempt_parts = after_attempt.split("_", 1)
                        if len(attempt_parts) > 1:
                            short_name = attempt_parts[1]
                            # Remove leading netid prefix if present
                            if "_" in short_name:
                                maybe_netid = short_name.split("_", 1)[0].lower()
                                if maybe_netid == sid.lower() or maybe_netid == sid:
                                    short_name = short_name.split("_", 1)[1]
                    file_links += f'<a href="{view_url}" target="_blank" class="file-link">{icon} {short_name}</a> <span style="color:#aaa;">({size_kb:.0f}KB)</span> '
                file_rows += f"<tr><td><code>{sid}</code></td><td>{file_links}</td></tr>"

            files_html = f"""
            <div class="card">
                <h2>Student Submissions ({len(students)} students)</h2>
                <table class="file-table">
                  <tr><th>Student ID</th><th>Files</th></tr>
                  {file_rows}
                </table>
            </div>"""

        if graded_files:
            graded_rows = ""
            for fn, ext, size_kb in graded_files:
                view_url = f"/api/submission-file?course={course_key}&assignment={assign_key}&file={urllib_parse.quote(fn)}&subdir=assignment2_graded"
                graded_rows += f'<tr><td><a href="{view_url}" target="_blank">{fn}</a></td><td>{size_kb:.0f}KB</td></tr>'
            files_html += f"""
            <div class="card">
                <h2>Graded Files</h2>
                <table class="file-table">
                  <tr><th>File</th><th>Size</th></tr>
                  {graded_rows}
                </table>
            </div>"""

    return base_html(f"Grading - {assign['label']}", f"""
    <div class="card">
      <h1>Assignment Grading</h1>
      <div>{course_btns}</div>
    </div>
    <div class="card">
      <h2>{course['label']} - {assign['label']}</h2>
      <div style="margin-bottom:1rem;">{assign_btns}</div>
      <p style="color:#666;">Folder: <code>{folder}</code> ({folder_info})</p>
      <div style="margin-top:1rem;">
        <button class="btn" id="gradeBtn"
                onclick="gradeAssignments('{course_key}', '{assign_key}')">
            Grade Assignments
        </button>
      </div>
      <div class="grading-progress" id="gradingLog"></div>
    </div>
    {results_html}
    {files_html}
    """, active="grading", extra_js=GRADING_JS)


@app.route("/api/submission-file")
def api_submission_file():
    from flask import send_file
    course_key = request.args.get("course", "")
    assign_key = request.args.get("assignment", "")
    filename = request.args.get("file", "")
    subdir = request.args.get("subdir", "")
    course = GRADING_COURSES.get(course_key, {})
    assign = course.get("assignments", {}).get(assign_key, {})
    if not assign or not filename:
        return "Not found", 404
    safe_name = os.path.basename(filename)
    folder = os.path.join(GRADING_DIR, assign["folder"])
    if subdir:
        folder = os.path.join(folder, os.path.basename(subdir))
    file_path = os.path.join(folder, safe_name)
    if not os.path.exists(file_path):
        return "File not found", 404
    ext = os.path.splitext(safe_name)[1].lower()
    if ext == ".pdf":
        return send_file(file_path, mimetype='application/pdf', download_name=safe_name)
    elif ext == ".sql":
        return send_file(file_path, mimetype='text/plain', download_name=safe_name)
    elif ext in (".docx", ".doc"):
        # Try to convert to PDF for inline viewing
        pdf_path = file_path.rsplit('.', 1)[0] + '.pdf'
        if not os.path.exists(pdf_path) or os.path.getmtime(file_path) > os.path.getmtime(pdf_path):
            try:
                pdf_path = docx_to_pdf(file_path)
            except Exception:
                return send_file(file_path, as_attachment=True, download_name=safe_name)
        return send_file(pdf_path, mimetype='application/pdf',
                         download_name=safe_name.rsplit('.', 1)[0] + '.pdf')
    elif ext == ".xlsx":
        return send_file(file_path, as_attachment=True, download_name=safe_name)
    return send_file(file_path, as_attachment=True, download_name=safe_name)


@app.route("/api/grade-start", methods=["POST"])
def api_grade_start():
    if grading_status["running"]:
        return jsonify({"error": "Grading is already running"})
    course_key = request.args.get("course", "")
    assign_key = request.args.get("assignment", "")
    t = threading.Thread(target=run_grading, args=(course_key, assign_key), daemon=True)
    t.start()
    return jsonify({"ok": True})


@app.route("/api/grade-status")
def api_grade_status():
    return jsonify(grading_status)


@app.route("/api/grade-download")
def api_grade_download():
    from flask import send_file
    course_key = request.args.get("course", "")
    assign_key = request.args.get("assignment", "")
    course = GRADING_COURSES.get(course_key, {})
    assign = course.get("assignments", {}).get(assign_key, {})
    if not assign:
        return "Not found", 404
    folder = os.path.join(GRADING_DIR, assign["folder"])
    excel_path = os.path.join(folder, f"{assign_key}_Grade_Upload.xlsx")
    if not os.path.exists(excel_path):
        return "Excel file not generated yet. Run grading first.", 404
    return send_file(excel_path, as_attachment=True,
                     download_name=f"{assign_key}_Grade_Upload.xlsx")


def docx_to_pdf(docx_path):
    """Convert a .docx file to PDF using MS Word COM automation. Returns PDF path."""
    import win32com.client
    import pythoncom
    pythoncom.CoInitialize()
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(docx_path))
        pdf_path = docx_path.rsplit('.', 1)[0] + '.pdf'
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
        word.Quit()
        return pdf_path
    finally:
        pythoncom.CoUninitialize()


@app.route("/api/graded-file")
def api_graded_file():
    from flask import send_file
    course_key = request.args.get("course", "")
    assign_key = request.args.get("assignment", "")
    filename = request.args.get("file", "")
    course = GRADING_COURSES.get(course_key, {})
    assign = course.get("assignments", {}).get(assign_key, {})
    if not assign or not filename:
        return "Not found", 404
    # Sanitize filename to prevent path traversal
    safe_name = os.path.basename(filename)
    folder = os.path.join(GRADING_DIR, assign["folder"])
    file_path = os.path.join(folder, safe_name)
    # Also check in the BUAN6320 subfolder where graded files are moved
    if not os.path.exists(file_path):
        file_path = os.path.join(folder, "BUAN6320", safe_name)
    if not os.path.exists(file_path):
        return "File not found", 404
    # Convert to PDF and serve inline
    pdf_path = file_path.rsplit('.', 1)[0] + '.pdf'
    if not os.path.exists(pdf_path) or os.path.getmtime(file_path) > os.path.getmtime(pdf_path):
        try:
            pdf_path = docx_to_pdf(file_path)
        except Exception as e:
            # Fallback to docx download if conversion fails
            return send_file(file_path, as_attachment=True, download_name=safe_name)
    pdf_name = safe_name.rsplit('.', 1)[0] + '.pdf'
    return send_file(pdf_path, mimetype='application/pdf', download_name=pdf_name)


@app.route("/grading/upload/<course_key>/<assign_key>", methods=["POST"])
def grading_upload(course_key, assign_key):
    course = GRADING_COURSES.get(course_key)
    assign = course["assignments"].get(assign_key) if course else None
    if not assign:
        flash("Invalid course/assignment.", "error")
        return redirect(url_for("grading_page"))

    folder = os.path.join(GRADING_DIR, assign["folder"])
    os.makedirs(folder, exist_ok=True)

    files = request.files.getlist("files")
    saved = 0
    for f in files:
        if f.filename and f.filename.endswith(".docx"):
            f.save(os.path.join(folder, f.filename))
            saved += 1

    flash(f"Uploaded {saved} files to {folder}", "success")
    return redirect(url_for("grading_assignment", course_key=course_key, assign_key=assign_key))


# ── Main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    init_db()

    conn = get_db()
    count = conn.execute("SELECT COUNT(*) FROM students").fetchone()[0]
    conn.close()

    if count == 0:
        roster_dir = os.path.expanduser("~/Downloads")
        roster_files = [
            "roster-buan4320.501.26s-2026-02-14-155649.xlsx",
            "roster-buan4320.502.26s-2026-02-14-155721.xlsx",
            "roster-buan4351.003.26s-2026-02-14-155826.xlsx",
            "roster-buan6320.s01.26s-2026-02-14-155846.xlsx",
        ]
        for rf in roster_files:
            path = os.path.join(roster_dir, rf)
            if os.path.exists(path):
                n = load_roster_from_path(path)
                print(f"  Loaded {n} students from {rf}")
            else:
                print(f"  WARNING: {rf} not found")

    conn = get_db()
    total = conn.execute("SELECT COUNT(*) FROM students").fetchone()[0]
    conn.close()
    print(f"\nTotal students in database: {total}")
    print(f"Starting on http://localhost:{PORT}")
    app.run(host="0.0.0.0", port=PORT, debug=False)
